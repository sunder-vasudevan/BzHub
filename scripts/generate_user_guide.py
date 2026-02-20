from docx import Document

guide = Document()
guide.add_heading('Inventory CRM - User Guide', level=1)

guide.add_heading('Quick Start', level=2)
guide.add_paragraph('Run from project folder:')
guide.add_paragraph('python inventory_crm_sqlite.py')
guide.add_paragraph('Default admin credentials:')
guide.add_paragraph('Username: admin  Password: admin123')

guide.add_heading('Login', level=2)
guide.add_paragraph('Enter your username and password on the login screen.\nRoles: admin has full access; other users have limited access.')

guide.add_heading('Inventory', level=2)
guide.add_paragraph('Add Item: All users can add new items via the Inventory tab.')
guide.add_paragraph('Update/Delete: Only admins can update or delete items; these controls are hidden for non-admins.')
guide.add_paragraph('Refresh: Click Refresh to reload inventory from the database.')

guide.add_heading('Point of Sale (POS)', level=2)
guide.add_paragraph('Add to Cart: Select an item, set quantity/price, click Add to Cart.')
guide.add_paragraph('Generate Bill: Click Generate Bill to create a bill preview.')
guide.add_paragraph('Preview: Bill opens in a preview window — review there.')
guide.add_paragraph('Print / Save: Use the preview\'s Print or Save As buttons to print or save the bill.')
guide.add_paragraph('Cart Controls: Remove or Clear Cart from the POS screen.')

guide.add_heading('Bills', level=2)
guide.add_paragraph('View Old Bills: Open the Bills tab to see previously generated bill_*.txt files.')
guide.add_paragraph('Preview/Print/Delete: Select a bill to preview, print, or delete it.')

guide.add_heading('Company & Email Settings', level=2)
guide.add_paragraph('Admin-only: Company Info and Email Settings tabs are only visible to admins. Admins can enter company details, save, and configure/test SMTP.')

guide.add_heading('Admin', level=2)
guide.add_paragraph('User Management: Admin tab lets you add/delete users and assign roles.')
guide.add_paragraph('Reports: Admins can generate Inventory, Sales, Profit and Activity Log reports.')

guide.add_heading('Troubleshooting', level=2)
guide.add_paragraph('Schema Errors: If you see errors like missing columns, delete inventory.db, inventory.db-shm, and inventory.db-wal then restart to recreate schema.')
guide.add_paragraph('No printer: If printing fails, save the bill and print from your OS file viewer.')
guide.add_paragraph('Logs: Activity log available in Reports → Activity Log.')

guide.add_heading('Notes & Next Steps', level=2)
guide.add_paragraph('Generated bills are saved as bill_YYYYMMDDHHMMSS.txt in the project folder.')
guide.add_paragraph('If you want the Bills tab to be admin-only or other UI changes, tell me and I\'ll update the app.')

outfile = 'Inventory_CRM_User_Guide.docx'
guide.save(outfile)
print('Created', outfile)
