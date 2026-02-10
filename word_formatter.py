import tkinter as tk
from tkinter import ttk, filedialog, colorchooser, messagebox
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class WordFormatterApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Word Document Formatter")
        self.root.geometry("600x800")
        self.doc = None
        self.selected_color = (0, 0, 0)  # Default black
        
        self.create_ui()
    
    def create_ui(self):
        # Frame for file operations
        file_frame = ttk.LabelFrame(self.root, text="File Operations", padding=10)
        file_frame.pack(fill="x", padx=10, pady=10)
        
        ttk.Button(file_frame, text="New Document", command=self.new_document).pack(side="left", padx=5)
        ttk.Button(file_frame, text="Open Document", command=self.open_document).pack(side="left", padx=5)
        ttk.Button(file_frame, text="Save Document", command=self.save_document).pack(side="left", padx=5)
        
        # Frame for text formatting
        format_frame = ttk.LabelFrame(self.root, text="Text Formatting", padding=10)
        format_frame.pack(fill="x", padx=10, pady=10)
        
        ttk.Label(format_frame, text="Font:").grid(row=0, column=0, sticky="w")
        self.font_var = tk.StringVar(value="Calibri")
        font_combo = ttk.Combobox(format_frame, textvariable=self.font_var, 
                                   values=["Calibri", "Arial", "Times New Roman", "Courier New", "Georgia"])
        font_combo.grid(row=0, column=1, sticky="ew", padx=5)
        
        ttk.Label(format_frame, text="Font Size:").grid(row=1, column=0, sticky="w")
        self.size_var = tk.StringVar(value="12")
        size_spin = ttk.Spinbox(format_frame, from_=8, to=72, textvariable=self.size_var)
        size_spin.grid(row=1, column=1, sticky="ew", padx=5)
        
        ttk.Label(format_frame, text="Color:").grid(row=2, column=0, sticky="w")
        ttk.Button(format_frame, text="Choose Color", command=self.choose_color).grid(row=2, column=1, sticky="ew", padx=5)
        self.color_label = ttk.Label(format_frame, text="■ Black", relief="solid", borderwidth=1)
        self.color_label.grid(row=2, column=2, padx=5)
        
        # Style checkboxes
        style_frame = ttk.Frame(format_frame)
        style_frame.grid(row=3, column=0, columnspan=3, sticky="ew", pady=10)
        
        self.bold_var = tk.BooleanVar()
        self.italic_var = tk.BooleanVar()
        self.underline_var = tk.BooleanVar()
        
        ttk.Checkbutton(style_frame, text="Bold", variable=self.bold_var).pack(side="left", padx=5)
        ttk.Checkbutton(style_frame, text="Italic", variable=self.italic_var).pack(side="left", padx=5)
        ttk.Checkbutton(style_frame, text="Underline", variable=self.underline_var).pack(side="left", padx=5)
        
        # Alignment
        ttk.Label(format_frame, text="Alignment:").grid(row=4, column=0, sticky="w")
        self.alignment_var = tk.StringVar(value="Left")
        alignment_combo = ttk.Combobox(format_frame, textvariable=self.alignment_var, 
                                        values=["Left", "Center", "Right", "Justify"], state="readonly")
        alignment_combo.grid(row=4, column=1, sticky="ew", padx=5)
        
        # Text input
        text_frame = ttk.LabelFrame(self.root, text="Add Content", padding=10)
        text_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        ttk.Label(text_frame, text="Content Type:").pack(anchor="w")
        self.content_type_var = tk.StringVar(value="Paragraph")
        type_combo = ttk.Combobox(text_frame, textvariable=self.content_type_var,
                                   values=["Paragraph", "Heading 1", "Heading 2", "Heading 3", "Bullet Point", "Table"], state="readonly")
        type_combo.pack(fill="x", pady=5)
        
        ttk.Label(text_frame, text="Text Content:").pack(anchor="w")
        self.text_input = tk.Text(text_frame, height=8, width=60)
        self.text_input.pack(fill="both", expand=True, pady=5)
        
        # Table options (initially hidden)
        self.table_frame = ttk.LabelFrame(self.root, text="Table Options", padding=10)
        
        ttk.Label(self.table_frame, text="Rows:").grid(row=0, column=0)
        self.table_rows = ttk.Spinbox(self.table_frame, from_=1, to=20, width=5)
        self.table_rows.set(3)
        self.table_rows.grid(row=0, column=1, sticky="w")
        
        ttk.Label(self.table_frame, text="Columns:").grid(row=1, column=0)
        self.table_cols = ttk.Spinbox(self.table_frame, from_=1, to=10, width=5)
        self.table_cols.set(3)
        self.table_cols.grid(row=1, column=1, sticky="w")
        
        # Buttons
        button_frame = ttk.Frame(self.root)
        button_frame.pack(fill="x", padx=10, pady=10)
        
        ttk.Button(button_frame, text="Add to Document", command=self.add_content).pack(side="left", padx=5)
        ttk.Button(button_frame, text="Clear Text", command=self.clear_text).pack(side="left", padx=5)
        
        format_frame.columnconfigure(1, weight=1)
        text_frame.columnconfigure(0, weight=1)
    
    def new_document(self):
        self.doc = Document()
        messagebox.showinfo("Success", "New document created!")
    
    def open_document(self):
        filepath = filedialog.askopenfilename(filetypes=[("Word files", "*.docx"), ("All files", "*.*")])
        if filepath:
            self.doc = Document(filepath)
            messagebox.showinfo("Success", f"Document opened: {filepath}")
    
    def save_document(self):
        if self.doc is None:
            messagebox.showerror("Error", "No document to save. Create or open a document first.")
            return
        
        filepath = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        if filepath:
            self.doc.save(filepath)
            messagebox.showinfo("Success", f"Document saved: {filepath}")
    
    def choose_color(self):
        color_tuple = colorchooser.askcolor(color="black")
        if color_tuple[0]:  # User didn't cancel
            self.selected_color = tuple(int(c) for c in color_tuple[0])
            self.color_label.config(text=f"■ RGB{self.selected_color}")
    
    def get_alignment(self):
        alignment_map = {
            "Left": WD_ALIGN_PARAGRAPH.LEFT,
            "Center": WD_ALIGN_PARAGRAPH.CENTER,
            "Right": WD_ALIGN_PARAGRAPH.RIGHT,
            "Justify": WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        return alignment_map.get(self.alignment_var.get(), WD_ALIGN_PARAGRAPH.LEFT)
    
    def add_content(self):
        if self.doc is None:
            messagebox.showerror("Error", "Create or open a document first.")
            return
        
        content_type = self.content_type_var.get()
        text = self.text_input.get("1.0", tk.END).strip()
        
        if not text and content_type != "Table":
            messagebox.showerror("Error", "Please enter text content.")
            return
        
        try:
            font_name = self.font_var.get()
            font_size = int(self.size_var.get())
            
            if content_type == "Paragraph":
                p = self.doc.add_paragraph(text)
                self._apply_formatting(p.runs[0] if p.runs else p.add_run(text), font_name, font_size)
                p.alignment = self.get_alignment()
            
            elif content_type.startswith("Heading"):
                level = int(content_type.split()[-1])
                p = self.doc.add_heading(text, level=level)
                self._apply_formatting(p.runs[0] if p.runs else p.add_run(text), font_name, font_size)
                p.alignment = self.get_alignment()
            
            elif content_type == "Bullet Point":
                p = self.doc.add_paragraph(text, style="List Bullet")
                self._apply_formatting(p.runs[0] if p.runs else p.add_run(text), font_name, font_size)
                p.alignment = self.get_alignment()
            
            elif content_type == "Table":
                rows = int(self.table_rows.get())
                cols = int(self.table_cols.get())
                table = self.doc.add_table(rows=rows, cols=cols)
                table.style = "Light Grid Accent 1"
                
                if text:
                    cells = table.rows[0].cells
                    for i, header in enumerate(text.split(",")):
                        if i < len(cells):
                            cells[i].text = header.strip()
            
            messagebox.showinfo("Success", f"{content_type} added to document!")
            self.text_input.delete("1.0", tk.END)
        
        except ValueError:
            messagebox.showerror("Error", "Invalid font size. Enter a number.")
    
    def _apply_formatting(self, run, font_name, font_size):
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = self.bold_var.get()
        run.font.italic = self.italic_var.get()
        run.font.underline = self.underline_var.get()
        run.font.color.rgb = RGBColor(*self.selected_color)
    
    def clear_text(self):
        self.text_input.delete("1.0", tk.END)

if __name__ == "__main__":
    root = tk.Tk()
    app = WordFormatterApp(root)
    root.mainloop()
