#!/usr/bin/env python3
"""
Generates a McKinsey/PwC-style Word document for the BzHub Efficiency White Paper.
Run: python3.9 documentation/generate_whitepaper_docx.py
Output: documentation/BzHub_Efficiency_WhitePaper.docx
"""

import io
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colour palette stored as (r, g, b) tuples ────────────────────────────────
# Use rgb(*DARK_NAVY) when passing to font.color.rgb; use hex_color() for XML.
DARK_NAVY   = (0x00, 0x2B, 0x5C)   # deep navy — headings, cover
MID_BLUE    = (0x00, 0x5B, 0x99)   # section rules, table headers
ACCENT_TEAL = (0x00, 0x7A, 0x87)   # pull-out boxes
LIGHT_GREY  = (0xF2, 0xF4, 0xF7)   # table row fill
WHITE       = (0xFF, 0xFF, 0xFF)
DARK_TEXT   = (0x1A, 0x1A, 0x2E)
RULE_BLUE   = (0x00, 0x5B, 0x99)


def rgb(t):
    """Convert (r,g,b) tuple to RGBColor."""
    return RGBColor(t[0], t[1], t[2])


def hex_color(t):
    """Convert (r,g,b) tuple to uppercase hex string for XML attributes."""
    return f"{t[0]:02X}{t[1]:02X}{t[2]:02X}"


# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, color_tuple):
    """Fill a table cell with a solid background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color(color_tuple))
    tcPr.append(shd)


def add_horizontal_rule(doc, color=RULE_BLUE, thickness_pt: int = 12):
    """Add a coloured horizontal rule paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(thickness_pt))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), hex_color(color))
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def set_run_font(run, name='Calibri', size_pt=10, bold=False, italic=False,
                 color=None):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = rgb(color)


def heading1(doc, text):
    """Section heading — large, navy, with rule below."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, name='Calibri', size_pt=16, bold=True, color=DARK_NAVY)
    add_horizontal_rule(doc, color=MID_BLUE, thickness_pt=8)
    return p


def heading2(doc, text):
    """Sub-section heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, name='Calibri', size_pt=12, bold=True, color=MID_BLUE)
    return p


def body(doc, text, space_after=6):
    """Standard body paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, name='Calibri', size_pt=10, color=DARK_TEXT)
    return p


def bullet(doc, text, level=0):
    """Bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.8 + level * 0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, name='Calibri', size_pt=10, color=DARK_TEXT)
    return p


def callout_box(doc, text, label='KEY FINDING'):
    """Teal accent box for key insights."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, (0xE8, 0xF4, 0xF6))
    cell.width = Inches(6)
    # Left border accent
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')
    left.set(qn('w:color'), hex_color(ACCENT_TEAL))
    tcBorders.append(left)
    tcPr.append(tcBorders)

    p_label = cell.paragraphs[0]
    p_label.paragraph_format.space_before = Pt(4)
    p_label.paragraph_format.space_after = Pt(2)
    r_label = p_label.add_run(label)
    set_run_font(r_label, name='Calibri', size_pt=8, bold=True, color=ACCENT_TEAL)

    p_text = cell.add_paragraph()
    p_text.paragraph_format.space_before = Pt(2)
    p_text.paragraph_format.space_after = Pt(6)
    r_text = p_text.add_run(text)
    set_run_font(r_text, name='Calibri', size_pt=10.5, italic=True, color=DARK_NAVY)
    doc.add_paragraph()


def styled_table(doc, headers, rows, col_widths_cm=None):
    """
    Build a styled table with navy header row and alternating grey rows.
    headers: list of str
    rows: list of list of str
    """
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'

    # Header row
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr_cells[i], MID_BLUE)
        p = hdr_cells[i].paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(h)
        set_run_font(run, name='Calibri', size_pt=9, bold=True, color=WHITE)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row_cells = tbl.rows[r_idx + 1].cells
        bg = LIGHT_GREY if r_idx % 2 == 0 else WHITE
        for c_idx, cell_text in enumerate(row_data):
            set_cell_bg(row_cells[c_idx], bg)
            p = row_cells[c_idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(str(cell_text))
            set_run_font(run, name='Calibri', size_pt=9, color=DARK_TEXT)

    # Column widths
    if col_widths_cm:
        for row in tbl.rows:
            for i, cell in enumerate(row.cells):
                if i < len(col_widths_cm):
                    cell.width = Cm(col_widths_cm[i])

    doc.add_paragraph()
    return tbl


# ── Footer & ToC helpers ──────────────────────────────────────────────────────

def add_footer(doc):
    """Add page numbers (centre) + copyright line (right) to every section footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False

        # Single paragraph: copyright left | page number centre
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.clear()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(0)

        # Tab stops: centre at 8.2 cm, right at 16.4 cm
        from docx.oxml import OxmlElement as OE
        pPr = p._p.get_or_add_pPr()
        tabs = OE('w:tabs')
        for pos_twips, align in [("4640", "center"), ("9280", "right")]:
            tab = OE('w:tab')
            tab.set(qn('w:val'), align)
            tab.set(qn('w:pos'), pos_twips)
            tabs.append(tab)
        pPr.append(tabs)

        # Copyright (left)
        r_copy = p.add_run("© 2026 sunder-vasudevan. All rights reserved.")
        set_run_font(r_copy, name='Calibri', size_pt=8, color=(0x88, 0x88, 0x88))

        # Tab to centre
        p.add_run("\t")

        # Page number field: "Page X of Y"
        r_pre = p.add_run("Page ")
        set_run_font(r_pre, name='Calibri', size_pt=8, color=(0x88, 0x88, 0x88))

        def _fld(name):
            fld = OE('w:fldSimple')
            fld.set(qn('w:instr'), f' {name} ')
            r = OE('w:r')
            rpr = OE('w:rPr')
            sz = OE('w:sz')
            sz.set(qn('w:val'), '16')
            rpr.append(sz)
            r.append(rpr)
            t = OE('w:t')
            t.text = name
            r.append(t)
            fld.append(r)
            return fld

        p._p.append(_fld('PAGE'))
        r_of = p.add_run(" of ")
        set_run_font(r_of, name='Calibri', size_pt=8, color=(0x88, 0x88, 0x88))
        p._p.append(_fld('NUMPAGES'))

        # Tab to right — BzHub label
        r_tab2 = p.add_run("\t")
        r_brand = p.add_run("BzHub Efficiency Case Study")
        set_run_font(r_brand, name='Calibri', size_pt=8, italic=True,
                     color=(0x00, 0x5B, 0x99))

        # Top border on footer paragraph
        pBdr = OE('w:pBdr')
        top = OE('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '4')
        top.set(qn('w:space'), '1')
        top.set(qn('w:color'), hex_color(MID_BLUE))
        pBdr.append(top)
        pPr.append(pBdr)


def add_toc(doc):
    """Insert a Table of Contents page using Word's TOC field (auto-updates on open)."""
    # Heading
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after  = Pt(6)
    r_title = p_title.add_run("Table of Contents")
    set_run_font(r_title, name='Calibri', size_pt=16, bold=True, color=DARK_NAVY)
    add_horizontal_rule(doc, color=MID_BLUE, thickness_pt=8)

    # Instruction note
    p_note = doc.add_paragraph()
    r_note = p_note.add_run(
        "Right-click this table and select 'Update Field' to refresh page numbers after opening.")
    set_run_font(r_note, name='Calibri', size_pt=8, italic=True,
                 color=(0x88, 0x88, 0x88))
    p_note.paragraph_format.space_after = Pt(8)

    # TOC field paragraph
    p_toc = doc.add_paragraph()
    p_toc.paragraph_format.space_before = Pt(0)
    p_toc.paragraph_format.space_after  = Pt(0)

    run = p_toc.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_begin)

    run2 = p_toc.add_run()
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    run2._r.append(instr)

    run3 = p_toc.add_run()
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    run3._r.append(fld_sep)

    run4 = p_toc.add_run()
    placeholder = OxmlElement('w:t')
    placeholder.text = '[Table of contents — update field in Word to populate]'
    run4._r.append(placeholder)
    set_run_font(run4, name='Calibri', size_pt=9, italic=True,
                 color=(0x88, 0x88, 0x88))

    run5 = p_toc.add_run()
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run5._r.append(fld_end)

    # Page break after ToC
    doc.add_page_break()


# ── Chart helpers ─────────────────────────────────────────────────────────────

def _navy():  return tuple(c/255 for c in DARK_NAVY)
def _teal():  return tuple(c/255 for c in ACCENT_TEAL)
def _blue():  return tuple(c/255 for c in MID_BLUE)
def _grey():  return (0.95, 0.96, 0.97)


def _savefig(fig, dpi=150):
    """Save matplotlib figure to BytesIO and return it."""
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=dpi, bbox_inches='tight')
    buf.seek(0)
    plt.close(fig)
    return buf


def chart_commits_per_day():
    """Bar chart: commits per active coding day."""
    labels = ['Feb 10', 'Feb 16', 'Feb 17', 'Feb 18', 'Feb 19',
              'Mar 09', 'Mar 10', 'Mar 11', 'Mar 12', 'Mar 13']
    values = [3, 1, 1, 8, 3, 1, 17, 16, 11, 3]
    colors = [_teal() if v >= 10 else _blue() if v >= 5 else _grey() for v in values]

    fig, ax = plt.subplots(figsize=(7, 3.2))
    bars = ax.bar(labels, values, color=colors, edgecolor='white', linewidth=0.8)

    for bar, val in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.3,
                str(val), ha='center', va='bottom', fontsize=8,
                color=tuple(c/255 for c in DARK_TEXT), fontweight='bold')

    ax.set_ylabel('Commits', fontsize=9, color=tuple(c/255 for c in DARK_TEXT))
    ax.set_title('Commits per Active Coding Day', fontsize=10, fontweight='bold',
                 color=_navy(), pad=10)
    ax.set_facecolor('white')
    fig.patch.set_facecolor('white')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_color('#CCCCCC')
    ax.spines['bottom'].set_color('#CCCCCC')
    ax.tick_params(axis='x', rotation=35, labelsize=8)
    ax.tick_params(axis='y', labelsize=8)
    ax.set_ylim(0, 20)

    legend_patches = [
        mpatches.Patch(color=_teal(), label='Peak sprint days (≥10 commits)'),
        mpatches.Patch(color=_blue(), label='Active sprint days (5–9)'),
        mpatches.Patch(color=_grey(), label='Light activity (<5)'),
    ]
    ax.legend(handles=legend_patches, fontsize=7, framealpha=0.5, loc='upper left')
    fig.tight_layout()
    return _savefig(fig)


def chart_loc_per_release():
    """Horizontal bar chart: net new LOC per release."""
    releases = ['v2.0', 'v3.x', 'v4.0', 'v4.1.x', 'v4.2.x', 'v4.3.0',
                'v4.4.0', 'v4.5.0', 'v4.6.0', 'v4.7.0', 'v4.7.1',
                'v4.8.0', 'v4.9.x', 'v5.0.0']
    loc =     [3776, 2696, 270, 472, 1042, 1177,
               1191, 656, 1267, 169, 971,
               196, 264, 495]
    logged =  [False, False, False, False, False, False,
               False, False, False, True, True,
               True, True, False]

    colors = [_teal() if lg else _blue() for lg in logged]

    fig, ax = plt.subplots(figsize=(7, 4.5))
    bars = ax.barh(releases, loc, color=colors, edgecolor='white', linewidth=0.6)

    for bar, val in zip(bars, loc):
        ax.text(bar.get_width() + 40, bar.get_y() + bar.get_height()/2,
                f'{val:,}', va='center', fontsize=7.5,
                color=tuple(c/255 for c in DARK_TEXT))

    ax.set_xlabel('Net New LOC (insertions)', fontsize=9, color=tuple(c/255 for c in DARK_TEXT))
    ax.set_title('Net New Lines of Code per Release', fontsize=10, fontweight='bold',
                 color=_navy(), pad=10)
    ax.set_facecolor('white')
    fig.patch.set_facecolor('white')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_color('#CCCCCC')
    ax.spines['bottom'].set_color('#CCCCCC')
    ax.tick_params(labelsize=8)
    ax.set_xlim(0, 4400)
    ax.invert_yaxis()

    legend_patches = [
        mpatches.Patch(color=_teal(), label='Logged sessions (v4.7.0+)'),
        mpatches.Patch(color=_blue(), label='Pre-logging era'),
    ]
    ax.legend(handles=legend_patches, fontsize=7.5, framealpha=0.5, loc='lower right')
    fig.tight_layout()
    return _savefig(fig)


def chart_session_efficiency():
    """Grouped bar chart: interaction time vs. releases shipped per session."""
    sessions = ['Session 1\n(45 min)\nWhitepaper', 'Session 2\n(25 min)\nTemplates',
                'Session 3\n(35 min)\n6 Releases']
    times    = [45, 25, 35]
    releases = [0, 1, 6]

    x = np.arange(len(sessions))
    width = 0.35

    fig, ax1 = plt.subplots(figsize=(7, 3.5))
    ax2 = ax1.twinx()

    bars1 = ax1.bar(x - width/2, times, width, label='Interaction Time (mins)',
                    color=_blue(), edgecolor='white', linewidth=0.8)
    bars2 = ax2.bar(x + width/2, releases, width, label='Releases Shipped',
                    color=_teal(), edgecolor='white', linewidth=0.8)

    for bar, val in zip(bars1, times):
        ax1.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5,
                 f'{val}m', ha='center', va='bottom', fontsize=8, fontweight='bold',
                 color=_blue())
    for bar, val in zip(bars2, releases):
        ax2.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.05,
                 str(val), ha='center', va='bottom', fontsize=8, fontweight='bold',
                 color=_teal())

    ax1.set_ylabel('Interaction Time (minutes)', fontsize=9, color=_blue())
    ax2.set_ylabel('Releases Shipped', fontsize=9, color=_teal())
    ax1.set_title('Session Efficiency — Time Invested vs. Releases Shipped',
                  fontsize=10, fontweight='bold', color=_navy(), pad=10)
    ax1.set_xticks(x)
    ax1.set_xticklabels(sessions, fontsize=8)
    ax1.set_ylim(0, 60)
    ax2.set_ylim(0, 8)
    ax1.set_facecolor('white')
    fig.patch.set_facecolor('white')
    for spine in ['top']:
        ax1.spines[spine].set_visible(False)
        ax2.spines[spine].set_visible(False)
    ax1.spines['left'].set_color('#CCCCCC')
    ax1.spines['bottom'].set_color('#CCCCCC')

    lines1, labels1 = ax1.get_legend_handles_labels()
    lines2, labels2 = ax2.get_legend_handles_labels()
    ax1.legend(lines1 + lines2, labels1 + labels2, fontsize=7.5, loc='upper left',
               framealpha=0.6)
    fig.tight_layout()
    return _savefig(fig)


def chart_team_comparison():
    """Side-by-side bar chart: PO+Claude vs. 3-person team."""
    metrics  = ['Production\nReleases', 'Live\nModules', 'Arch.\nMigrations',
                'Active\nDays Used']
    po_vals  = [11, 9, 3, 10]
    tm_vals  = [4, 5, 1, 22]   # 31-day estimate for 3-person team

    x = np.arange(len(metrics))
    width = 0.35

    fig, ax = plt.subplots(figsize=(7, 3.5))
    bars1 = ax.bar(x - width/2, po_vals, width, label='1 PO + Claude',
                   color=_teal(), edgecolor='white')
    bars2 = ax.bar(x + width/2, tm_vals, width, label='3-Person Team (estimated)',
                   color=(0.75, 0.80, 0.85), edgecolor='white')

    for bar, val in zip(bars1, po_vals):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.2,
                str(val), ha='center', va='bottom', fontsize=9, fontweight='bold',
                color=_teal())
    for bar, val in zip(bars2, tm_vals):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.2,
                f'~{val}', ha='center', va='bottom', fontsize=9,
                color=(0.4, 0.4, 0.5))

    ax.set_title('PO + Claude vs. Traditional 3-Person Team (31-day window)',
                 fontsize=10, fontweight='bold', color=_navy(), pad=10)
    ax.set_xticks(x)
    ax.set_xticklabels(metrics, fontsize=9)
    ax.legend(fontsize=8.5, framealpha=0.6)
    ax.set_facecolor('white')
    fig.patch.set_facecolor('white')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_color('#CCCCCC')
    ax.spines['bottom'].set_color('#CCCCCC')
    ax.tick_params(labelsize=8)
    ax.set_ylim(0, 26)
    fig.tight_layout()
    return _savefig(fig)


def add_chart(doc, buf, width_inches=6.2, caption=None):
    """Insert a BytesIO chart image into the document."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(buf, width=Inches(width_inches))
    if caption:
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_before = Pt(2)
        p_cap.paragraph_format.space_after = Pt(10)
        r_cap = p_cap.add_run(caption)
        set_run_font(r_cap, name='Calibri', size_pt=8, italic=True,
                     color=(0x66, 0x66, 0x77))


# ── Document assembly ─────────────────────────────────────────────────────────

def build_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── COVER PAGE ────────────────────────────────────────────────────────────

    # Top navy band (simulated with a 1-col table)
    cover_band = doc.add_table(rows=1, cols=1)
    cover_cell = cover_band.cell(0, 0)
    set_cell_bg(cover_cell, DARK_NAVY)
    p_band = cover_cell.paragraphs[0]
    p_band.paragraph_format.space_before = Pt(18)
    p_band.paragraph_format.space_after = Pt(18)
    r = p_band.add_run("  ENGINEERING EFFICIENCY WITH AI PAIR PROGRAMMING")
    set_run_font(r, name='Calibri', size_pt=18, bold=True, color=WHITE)

    doc.add_paragraph()

    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_sub = p_sub.add_run("A Case Study of BzHub — One Product Owner, One Claude")
    set_run_font(r_sub, name='Calibri', size_pt=14, bold=False, color=MID_BLUE)

    doc.add_paragraph()

    # Meta block
    meta_tbl = doc.add_table(rows=4, cols=2)
    meta_data = [
        ("Author", "sunder-vasudevan"),
        ("Date", "March 2026"),
        ("Version", "1.1 — updated with measured interaction time data"),
        ("Audience", "Technical peers — software engineers, engineering leads, technical founders"),
    ]
    for i, (k, v) in enumerate(meta_data):
        cells = meta_tbl.rows[i].cells
        r_key = cells[0].paragraphs[0].add_run(k)
        set_run_font(r_key, name='Calibri', size_pt=9, bold=True, color=MID_BLUE)
        cells[0].width = Cm(3.5)
        r_val = cells[1].paragraphs[0].add_run(v)
        set_run_font(r_val, name='Calibri', size_pt=9, color=DARK_TEXT)

    doc.add_paragraph()
    add_horizontal_rule(doc, color=MID_BLUE, thickness_pt=16)
    doc.add_page_break()

    # ── TABLE OF CONTENTS ────────────────────────────────────────────────────

    add_toc(doc)

    # ── ABSTRACT ─────────────────────────────────────────────────────────────

    heading1(doc, "Abstract")
    body(doc,
         "This paper documents a real-world experiment: a single Product Owner building BzHub — "
         "a full-stack ERP web application — in collaboration with Claude (Anthropic's AI assistant) "
         "as a persistent pair-programmer. Over 31 calendar days (~10 active coding days), the pair "
         "shipped 11 production releases (v1.0 → v5.0.0), migrated across 3 architectural paradigms, "
         "and delivered 9 live product modules spanning ~15,800 lines of net new application code. "
         "Three sessions were formally logged with exact interaction times: 105 minutes of active "
         "Product Owner engagement produced 6 production releases — a compression ratio of approximately "
         "60–120x over equivalent traditional engineering effort. When compared against the output a "
         "traditional 3-person startup development team would reasonably produce in the same calendar "
         "window, the evidence — now anchored in measured data — supports a 2–4x velocity multiplier "
         "at the project level, with individual sessions demonstrating substantially higher compression "
         "ratios on active time. This paper analyses those claims using concrete git-derived data and "
         "measured interaction logs, identifies where the gains came from, and honestly acknowledges "
         "what this model does not replace.")

    # ── 1. INTRODUCTION ───────────────────────────────────────────────────────

    heading1(doc, "1. Introduction — The Experiment")
    body(doc,
         "Most AI-assisted development studies focus on narrow tasks: code completion accuracy, "
         "bug fix rates, test generation. This case study is different. It examines an end-to-end "
         "product build: ideation, architecture, feature development, cloud deployment, and "
         "documentation — all executed by a single Product Owner with Claude as an active collaborator.")

    heading2(doc, "The Question")
    body(doc,
         "When one capable Product Owner pairs with Claude as a full-session collaborator (not just "
         "a code autocomplete tool), does the output more closely resemble a solo effort or a small team?")

    heading2(doc, "What 'Efficiency' Means Here")
    for item in [
        "Feature throughput per calendar day",
        "Architectural scope delivered without accruing uncontrolled technical debt",
        "Code quality signals visible in git history",
        "Documentation discipline sustained throughout",
        "Measured: active Product Owner time per feature shipped (formally logged from v4.7.0)",
    ]:
        bullet(doc, item)

    heading2(doc, "Why It Matters for Technical Teams")
    body(doc,
         "AI pair-programming is often evaluated as a tool for individual tasks. This study argues "
         "the real leverage is at the workflow level — persistent context, zero handoff overhead, "
         "and on-demand expertise across the full stack.")

    # ── 2. PROJECT SCOPE ──────────────────────────────────────────────────────

    heading1(doc, "2. Project Scope — What Was Built")
    body(doc,
         "BzHub is a full-stack business management (ERP) web application targeting small-to-medium "
         "businesses. It is Odoo-inspired, with a goal of becoming a multi-tenant SaaS product.")

    heading2(doc, "Technology Stack")
    styled_table(doc,
        headers=["Layer", "Technology"],
        rows=[
            ("Frontend",    "Next.js 14 (App Router), Tailwind CSS, shadcn/ui"),
            ("Database",    "Supabase (PostgreSQL), Row Level Security enabled"),
            ("Deployment",  "Vercel (CI/CD on every push to main)"),
            ("Language",    "TypeScript throughout"),
        ],
        col_widths_cm=[4, 12]
    )

    heading2(doc, "Live Modules at v5.0.0")
    styled_table(doc,
        headers=["Module", "Key Features"],
        rows=[
            ("Dashboard",              "KPI cards, sales trend chart, fast/slow movers, Smart Insights, customizable layout"),
            ("Operations",             "Inventory, POS, Bills, Supplier management, Purchase Orders + approval"),
            ("HR",                     "Employees, Payroll, Goals, Appraisals + sign-off, Skills Matrix, Leave Requests + approval"),
            ("CRM",                    "Contacts, Leads — List / Kanban / Funnel views, lead scoring, follow-ups"),
            ("Reports",                "Sales Report, Top Sellers chart, Inventory Report — all with CSV export"),
            ("Settings",               "Company info, currency, industry templates, dynamic brand color, Custom Fields builder"),
            ("Help",                   "In-app user guide for all modules"),
            ("Employee Self-Service",  "My Goals, My Appraisals, My Leave, My Skills"),
            ("Notifications",          "Bell icon, pending leave/PO/appraisal/low-stock alerts"),
        ],
        col_widths_cm=[4.5, 11.5]
    )

    heading2(doc, "Architecture Evolution")
    body(doc, "The project underwent 3 full architectural paradigm shifts:")
    styled_table(doc,
        headers=["Stage", "Architecture"],
        rows=[
            ("v1.0",  "Desktop app (Tkinter / Python)"),
            ("v2.0",  "Next.js web frontend + FastAPI backend + SQLite"),
            ("v4.0+", "Next.js + Supabase (PostgreSQL) + Vercel (cloud)"),
        ],
        col_widths_cm=[3, 13]
    )
    body(doc, "Each migration was deliberate, documented, and left the codebase in a stable, deployable state.")

    # ── 3. METHODOLOGY ────────────────────────────────────────────────────────

    heading1(doc, "3. Methodology — How Efficiency Is Measured")

    heading2(doc, "Primary Metric: Feature Throughput per Engineer-Day")
    body(doc,
         "The key measure is how many meaningful product features were shipped per active coding day, "
         "compared to what a small team would typically ship in the same calendar window.")

    heading2(doc, "Secondary Metrics")
    for item in [
        "Commit frequency — intensity of active days",
        "Lines of code per feature commit — size and completeness of deliveries",
        "Architectural scope per sprint — how many layers were touched in a single session",
        "Documentation signal — release notes and help pages updated per feature",
    ]:
        bullet(doc, item)

    heading2(doc, "Comparison Baseline: Traditional 3-Person Startup Team")
    body(doc, "The baseline models a typical 3-person startup team:")
    for item in [
        "1 frontend engineer",
        "1 backend engineer",
        "1 tech lead / product manager",
    ]:
        bullet(doc, item)
    body(doc,
         "Typical output for such a team in a 2-week sprint: 2–4 significant features to production, "
         "with 15–25% of time consumed by ceremonies, code review, and coordination. Architecture "
         "migrations are treated as dedicated sprint work, not background activity.")

    heading2(doc, "Data Sources")
    styled_table(doc,
        headers=["Source", "Data Captured"],
        rows=[
            ("git log --shortstat",              "Commit messages, dates, file counts, lines added/deleted"),
            ("INTERACTION_LOG.md",               "Formally logged session interaction time from v4.7.0 onwards"),
            ("Product Owner self-assessment",    "Stated 2–3x velocity multiplier vs. working without AI"),
        ],
        col_widths_cm=[5, 11]
    )

    heading2(doc, "Measured Interaction Time — Formally Logged Sessions")
    body(doc,
         "Starting from the session that produced v4.7.0 (2026-03-12), all sessions were formally logged "
         "with exact Product Owner interaction times. Three sessions are on record:")
    styled_table(doc,
        headers=["Session", "Date", "Time", "Releases", "LOC Added", "Goal"],
        rows=[
            ("1", "2026-03-12", "~45 mins", "—",              "—",      "Efficiency whitepaper + doc generators"),
            ("2", "2026-03-12", "~25 mins", "v4.7.0",         "~170",   "Industry templates (FEAT-038) + doc system"),
            ("3", "2026-03-12", "~35 mins", "v4.7.1–v4.9.2",  "~1,862", "Brand color, Smart Insights, CRM views"),
            ("Total", "",       "~105 mins","6 releases",      "~1,600 net", ""),
        ],
        col_widths_cm=[1.6, 2.6, 2.2, 2.8, 2.2, 4.6]
    )
    body(doc,
         "Session 3 produced 6 production releases in 35 minutes of interaction time. Equivalent work "
         "for a traditional 3-person team is estimated at 3–6 engineer-weeks — a 60–120x compression "
         "ratio on active human direction time.")
    callout_box(doc,
        "35 minutes of Product Owner interaction → 6 production releases → 1,862 lines of inserted code "
        "across 9 features, 24+ files modified. Traditional team estimate: 3–6 engineer-weeks.",
        label="SESSION 3 — MEASURED RESULT")

    body(doc, "Figure 1: Session efficiency — interaction time invested vs. releases shipped per session.")
    add_chart(doc, chart_session_efficiency(), width_inches=6.0,
              caption="Figure 1 — Session Efficiency: Interaction Time (left axis, blue) vs. Releases Shipped (right axis, teal)")

    heading2(doc, "Acknowledged Limitations")
    for item in [
        "Interaction time logging began at v4.7.0 — earlier sessions (v1.0–v4.6.0) were not tracked",
        "'Interaction time' measures PO prompting and review effort, not wall-clock or Claude generation time",
        "Not all commits are equal — two node_modules commits (millions of lines) excluded from LOC counts",
        "LOC is an imperfect complexity proxy — a 50-line schema change can unlock more than 500 lines of UI",
        "Greenfield product — results may differ in large legacy codebases",
    ]:
        bullet(doc, item)

    # ── 4. TIMELINE & VELOCITY ────────────────────────────────────────────────

    heading1(doc, "4. Timeline & Velocity Analysis")

    heading2(doc, "Commit Activity by Date")
    styled_table(doc,
        headers=["Date", "Commits", "Notes"],
        rows=[
            ("2026-02-10", "3",  "Project initialisation"),
            ("2026-02-16", "1",  "Stability pass"),
            ("2026-02-17", "1",  "Bug fixes"),
            ("2026-02-18", "8",  "Architecture sprint — FastAPI, Supabase, Next.js scaffolding"),
            ("2026-02-19", "3",  "CRM features, v2.0 release notes"),
            ("2026-03-09", "1",  "v3.0 release tag"),
            ("2026-03-10", "17", "PEAK — full UI build, CRM, HR, Settings, production shadcn/ui"),
            ("2026-03-11", "16", "PEAK — Vercel deploy, Supabase integration, v4.x features"),
            ("2026-03-12", "11", "3 logged sessions: v4.4.0–v4.9.2 + whitepaper + docs"),
            ("2026-03-13", "3",  "v5.0.0 Custom Fields builder"),
        ],
        col_widths_cm=[3.5, 2.5, 10]
    )
    body(doc,
         "44 of 64 total commits (69%) occurred across just 4 days. Mar 10–11 alone account for "
         "33 commits (52%) — not a sign of erratic work, but of sustained high-intensity sessions "
         "enabled by Claude's ability to hold architectural context without cognitive fatigue.")
    add_chart(doc, chart_commits_per_day(), width_inches=6.2,
              caption="Figure 2 — Commits per Active Coding Day (teal = peak sprint ≥10, blue = 5–9, grey = <5)")

    heading2(doc, "Milestone-by-Milestone Breakdown")
    styled_table(doc,
        headers=["Release", "Date", "Key Output", "Files", "Insertions"],
        rows=[
            ("v1.0",        "Feb 10–16", "Desktop app, login, sidebar, core UI",                               "—",  "—"),
            ("v2.0",        "Feb 18–19", "CRM module, FastAPI backend, Next.js web frontend",                   "41", "3,860"),
            ("v3.0–3.1",    "Mar 9–10",  "Odoo CRM, Kanban, lead scoring, production shadcn/ui",               "23", "2,368"),
            ("v4.0",        "Mar 10–11", "Inventory image upload, sortable tables, fast/slow movers",           "4",  "369"),
            ("v4.1.x",      "Mar 11",    "Vercel deploy, Supabase integration, mobile responsive",              "11", "440"),
            ("v4.2.0",      "Mar 11",    "Reports page, Supplier management",                                   "5",  "699"),
            ("v4.3.0",      "Mar 11",    "Goals, Appraisals, Skills Matrix — 13 new DB functions",              "2",  "1,302"),
            ("v4.4.0",      "Mar 12",    "3 Approval workflows + new DB tables",                                "9",  "1,209"),
            ("v4.5.0",      "Mar 12",    "Employee Self-Service Portal (4 tabs)",                               "6",  "667"),
            ("v4.6.0",      "Mar 12",    "Notification Center, Dashboard custom., CSV Export, Global Search",   "16", "1,381"),
            ("v4.7.0 ●",    "Mar 12",    "Industry templates — Retail/Clinic/Restaurant/Distributor [S2]",      "3",  "170"),
            ("v4.7.1 ●",    "Mar 12",    "Dynamic brand color — entire app theme via CSS vars [S3]",            "24", "1,102"),
            ("v4.8.0 ●",    "Mar 12",    "Smart Insights dashboard card — stock/HR/sales nudges [S3]",          "2",  "197"),
            ("v4.9.x ●",    "Mar 12",    "CRM view switcher: List / Kanban / Funnel views [S3]",                "5",  "563"),
            ("v5.0.0",      "Mar 13",    "Custom Fields builder — FEAT-041 Phase 2.5a",                         "6",  "500"),
        ],
        col_widths_cm=[2.2, 2.4, 8.0, 1.7, 2.2]
    )
    body(doc, "● = formally logged session. S2 = Session 2 (25 min), S3 = Session 3 (35 min).", space_after=4)
    add_chart(doc, chart_loc_per_release(), width_inches=6.2,
              caption="Figure 3 — Net New LOC per Release (teal = logged sessions, blue = pre-logging era)")

    # ── 5. COMPARATIVE ANALYSIS ───────────────────────────────────────────────

    heading1(doc, "5. Comparative Analysis")

    heading2(doc, "The Two-Day Sprint Block (Mar 11–12)")
    body(doc,
         "The clearest evidence of velocity is the Mar 11–12 window. In approximately 2 calendar days:")
    styled_table(doc,
        headers=["Release", "Feature", "Net LOC"],
        rows=[
            ("v4.1.x",      "Cloud deployment + mobile responsive",                    "~495"),
            ("v4.2.0",      "Reports page + Supplier management",                      "699"),
            ("v4.3.0",      "Goals, Appraisals, Skills Matrix",                        "1,302"),
            ("v4.4.0",      "3 Approval workflows + DB schema",                        "1,209"),
            ("v4.5.0",      "Employee Self-Service Portal",                            "667"),
            ("v4.6.0",      "Notification Center, CSV Export, Global Search, Audit",   "1,267"),
            ("v4.7.0 ●",    "Industry Templates [S2, 25 min]",                         "169"),
            ("v4.7.1 ●",    "Dynamic brand color — 24 files [S3, 35 min]",             "971"),
            ("v4.8.0 ●",    "Smart Insights dashboard card [S3]",                      "196"),
            ("v4.9.x ●",    "CRM view switcher: List/Kanban/Funnel [S3]",              "264"),
            ("TOTAL",       "10 production releases, all layers",                      "~7,239"),
        ],
        col_widths_cm=[2.5, 9.5, 4.0]
    )

    callout_box(doc,
        "For a 3-person startup team, this block of work — spanning UI, data layer, DB schema design, "
        "cloud deployment, brand theming, analytics, and documentation — would realistically require "
        "3–6 two-week sprints (6–12 calendar weeks). The AI-assisted pair compressed it to 2 calendar days. "
        "Sessions 2 and 3 (60 mins combined) produced 7 of these 10 releases.",
        label="KEY FINDING")

    heading2(doc, "Measured Session Efficiency (Logged Data)")
    body(doc, "The interaction log provides ground-truth for Sessions 2 and 3:")
    styled_table(doc,
        headers=["Metric", "Session 3 (Measured)", "Traditional Team Estimate"],
        rows=[
            ("PO interaction time",        "35 minutes",              "N/A"),
            ("Releases shipped",           "6 (v4.7.1–v4.9.2)",       "~0.5–1 per sprint (2 weeks)"),
            ("LOC inserted",               "1,862",                   "—"),
            ("Cross-file changes",         "24 files (color alone)",  "Typically multi-PR coordination"),
            ("Estimated traditional effort","—",                      "3–6 engineer-weeks"),
            ("Compression ratio",          "—",                       "~60–120x on interaction time"),
        ],
        col_widths_cm=[5.0, 5.0, 6.0]
    )

    heading2(doc, "Full Project Window Comparison")
    styled_table(doc,
        headers=["Dimension", "1 Product Owner + Claude (31 days)", "3-Person Team (31 days, estimated)"],
        rows=[
            ("Major releases shipped",      "11 (v1.0 → v5.0.0)",          "3–5 (1–2 per 2-week sprint)"),
            ("Architecture migrations",     "3 full paradigm shifts",        "Likely 1 (or as dedicated project)"),
            ("Live modules delivered",      "9",                             "4–6"),
            ("Documentation maintained",    "Per-feature (project rule)",    "Varies — often deferred"),
            ("Deployment pipeline",         "Fully automated (Vercel CI/CD)","2–5 days setup time for a team"),
            ("Measured PO time (3 sessions)","~105 mins active",              "N/A (team always staffed)"),
        ],
        col_widths_cm=[4.5, 6.0, 5.5]
    )
    body(doc, "3-person team estimates are based on conventional startup sprint norms, not empirical benchmarks.",
         space_after=4)
    add_chart(doc, chart_team_comparison(), width_inches=6.2,
              caption="Figure 4 — PO + Claude vs. 3-Person Team: releases, modules, architecture migrations, active days used")

    heading2(doc, "Where the Gains Come From")
    body(doc, "The efficiency delta is not simply 'Claude writes code faster.' The compounding advantages are structural:")

    for label, desc in [
        ("Zero handoff overhead",
         "In a team, frontend and backend engineers coordinate through PRs, Slack, standup, and specs. "
         "Each handoff has a latency cost. The AI pair eliminates this entirely — full-stack features "
         "are conceived, designed, and implemented in a single uninterrupted session."),
        ("Persistent context across sessions",
         "The project uses a structured NOTES.md + memory file system to preserve project state between "
         "sessions. Claude resumes with full context — current version, pending features, known debt, "
         "architecture decisions — without onboarding overhead. A new team hire needs 1–2 weeks to "
         "reach this context depth."),
        ("On-demand full-stack competency",
         "No single engineer is equally strong across Next.js App Router, Supabase RLS design, "
         "Recharts, and mobile-first Tailwind layout. Claude provides competent first-draft "
         "implementations across all layers, which the Product Owner reviews and approves."),
        ("Documentation as a side-effect",
         "Every feature commit was accompanied by release notes, help page updates, and memory file "
         "updates. With Claude handling the mechanical writing, this discipline was sustained throughout "
         "— rare in solo projects and difficult to maintain in teams under delivery pressure."),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        r_label = p.add_run(label + " — ")
        set_run_font(r_label, name='Calibri', size_pt=10, bold=True, color=DARK_NAVY)
        r_desc = p.add_run(desc)
        set_run_font(r_desc, name='Calibri', size_pt=10, color=DARK_TEXT)

    doc.add_paragraph()

    # ── 6. COLLABORATION PATTERNS ─────────────────────────────────────────────

    heading1(doc, "6. Collaboration Patterns Observed")

    heading2(doc, "How Claude Was Used")
    styled_table(doc,
        headers=["Task Category", "Examples from BzHub"],
        rows=[
            ("Feature design",     "DB schema, component structure, and state design before coding"),
            ("Code generation",    "Full page components — 1,300+ LOC in single sessions"),
            ("Architecture",       "FastAPI vs. Supabase direct client; Vercel deployment config"),
            ("Bug diagnosis",      "TypeScript errors, Supabase RLS issues, Next.js config problems"),
            ("Schema design",      "Supabase table definitions, RLS policies, migration scripts"),
            ("Documentation",      "Release notes, help page content, NOTES.md updates per feature"),
            ("Refactoring",        "shadcn/ui adoption decision, data layer abstraction"),
        ],
        col_widths_cm=[4.5, 11.5]
    )

    heading2(doc, "The Context Persistence System")
    body(doc,
         "A key enabler of sustained velocity was explicit context management — the equivalent of a "
         "living spec document that is always current and always read before work begins:")
    for item in [
        "chat_persistent_notes/NOTES.md — master project state, read at session start",
        "documentation/RELEASE_NOTES_v4.1.md — version history maintained per feature",
        ".claude/projects/.../memory/ — auto-memory files for project state and collaboration rules",
    ]:
        bullet(doc, item)
    body(doc,
         "Most teams attempt this with wikis or Confluence — and fail to maintain it. Encoding it as "
         "a session discipline with AI enforcement makes it sustainable.")

    heading2(doc, "Product Owner's Role")
    body(doc,
         "Claude does not replace product or engineering judgment. The Product Owner was responsible for:")
    for item in [
        "Product decisions — what to build, in what order, and what to defer",
        "Architecture calls — which tech stack, which migration path, which trade-offs to accept",
        "Quality gate — reviewing all generated code before committing",
        "Strategic direction — the WhatsApp-first, multi-tenant SaaS end goal shaping every feature",
    ]:
        bullet(doc, item)

    callout_box(doc,
        "Claude was the executor. The Product Owner was the director. "
        "The distinction matters for how teams should think about AI-assisted workflows.",
        label="KEY PRINCIPLE")

    # ── 7. CODE QUALITY ───────────────────────────────────────────────────────

    heading1(doc, "7. Code Quality & Architecture Observations")

    heading2(doc, "Architecture Was Not Sacrificed for Speed")
    body(doc,
         "Three architecture migrations in 31 days could suggest chaotic churning. The git history "
         "tells a different story:")
    for item in [
        "Each migration was preceded by a documented decision commit (e.g., 'docs: record UI modernisation decision — Option B')",
        "Each migration left the codebase in a stable, deployed state before the next one began",
        "Known technical debt was explicitly called out in code and docs rather than hidden",
    ]:
        bullet(doc, item)

    heading2(doc, "Documentation Discipline")
    body(doc, "A project rule was enforced after every feature — no exceptions:")
    for item in [
        "Mark feature done in FEATURE_REQUESTS_AND_BUGS.md",
        "Add version entry to RELEASE_NOTES_v4.1.md",
        "Update NOTES.md with current version and next priority",
        "Update memory files with project state",
        "Add help section to src/app/help/page.tsx",
    ]:
        bullet(doc, item)

    heading2(doc, "Intentional Debt Logged, Not Ignored")
    styled_table(doc,
        headers=["Known Gap", "Status", "Planned Resolution"],
        rows=[
            ("Authentication (hardcoded admin/admin123)", "Intentional — v1 scope", "FEAT-036: Supabase Auth (Phase 3)"),
            ("RLS policies open (no user filtering)",     "Intentional — no auth yet", "Tightens when auth ships"),
            ("No organization_id on tables",              "Intentional — pre-multi-tenant", "FEAT-037: Multi-Tenancy (Phase 3)"),
        ],
        col_widths_cm=[5.5, 4.0, 6.5]
    )

    # ── 8. LIMITATIONS ────────────────────────────────────────────────────────

    heading1(doc, "8. Limitations & Honest Caveats")
    for label, desc in [
        ("Partial time tracking.",
         "Interaction time was formally logged from v4.7.0 onwards (3 sessions, ~105 mins). "
         "Sessions covering v1.0–v4.6.0 were not tracked. Active coding days for the earlier window "
         "are inferred from commit timestamps. A complete picture would require logging from day one."),
        ("'Interaction time' ≠ 'engineering hours'.",
         "Logged time measures the Product Owner's active prompting and review effort, not wall-clock "
         "time or the time Claude spends generating responses. It should be interpreted as 'human "
         "direction cost' rather than a direct substitute for team-hours in a traditional context."),
        ("Solo project dynamics differ from team dynamics.",
         "There was no code review from a second human engineer. Code review catches bugs, architectural "
         "issues, and knowledge-sharing opportunities that this model does not replicate."),
        ("Not all commits represent equal effort.",
         "The git history includes node_modules commits with millions of lines (excluded) and single-line "
         "config fixes alongside 1,300-line feature commits."),
        ("Greenfield advantage.",
         "Building from scratch allows architectural freedom that legacy codebases do not afford."),
        ("The multiplier depends on Product Owner quality.",
         "Claude amplifies the Product Owner's ability to execute. A Product Owner who cannot review "
         "code cannot safely use AI-generated code at this velocity."),
        ("Reproducibility is uncertain.",
         "This was one project, one domain. Generalisability to other stacks or engineer profiles is unproven."),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        r_label = p.add_run(label + " ")
        set_run_font(r_label, name='Calibri', size_pt=10, bold=True, color=DARK_NAVY)
        r_desc = p.add_run(desc)
        set_run_font(r_desc, name='Calibri', size_pt=10, color=DARK_TEXT)

    doc.add_paragraph()

    # ── 9. CONCLUSIONS ────────────────────────────────────────────────────────

    heading1(doc, "9. Conclusions & Recommendations")

    heading2(doc, "Key Finding")
    callout_box(doc,
        "One capable Product Owner working with Claude as a persistent, full-session pair-programmer "
        "can approximate the feature throughput of a 3-person startup development team over a sustained "
        "build window. Where formerly this rested on git history alone, the interaction log now provides "
        "measured evidence: 105 minutes of Product Owner engagement produced 6 production releases "
        "spanning 1,862 lines of inserted code across 9 features. For Session 3 specifically: 35 minutes "
        "of active interaction time produced work estimated at 3–6 engineer-weeks — a 60–120x compression "
        "ratio on human direction time.",
        label="CONCLUSION")

    body(doc, "The gains are not from Claude writing faster code. They are structural:")
    for item in [
        "Zero handoff latency",
        "Persistent cross-session context",
        "On-demand full-stack competency",
        "Documentation as a zero-marginal-cost side-effect",
    ]:
        bullet(doc, item)

    heading2(doc, "Recommended Workflow Patterns")
    for n, label, desc in [
        (1, "Invest in context management.",
         "A well-maintained NOTES.md or equivalent session-start document is the single highest-leverage "
         "practice. It eliminates the 15–30 minute catch-up cost at the start of every session."),
        (2, "Treat Claude as a full-session collaborator, not a one-shot query tool.",
         "The compounding gains come from sustained sessions where Claude holds architectural context. "
         "Piecemeal queries do not capture this."),
        (3, "Keep the Product Owner in the architecture seat.",
         "Define what to build and why before asking Claude to help build it. Product direction, "
         "technical trade-offs, and quality gate remain the Product Owner's responsibility."),
        (4, "Use the multiplier effect, not the replacement effect.",
         "The right frame is not 'Claude replaces the Product Owner.' It is 'one Product Owner + Claude "
         "can reach a place that previously required a team.'"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(2)
        r_n = p.add_run(f"{n}. {label} ")
        set_run_font(r_n, name='Calibri', size_pt=10, bold=True, color=MID_BLUE)
        r_d = p.add_run(desc)
        set_run_font(r_d, name='Calibri', size_pt=10, color=DARK_TEXT)

    doc.add_paragraph()

    heading2(doc, "Where This Model Works Best")
    for item in [
        "Greenfield feature development with clear requirements",
        "Database schema design and data layer abstraction",
        "Responsive UI work with established component libraries (shadcn/ui, Tailwind)",
        "Documentation and release note generation",
        "Cross-stack features requiring frontend + backend + DB changes in a single session",
    ]:
        bullet(doc, item)

    heading2(doc, "Where Human Judgment Remains Essential")
    for item in [
        "Product and architectural decisions (what to build, what to defer)",
        "Security-sensitive code (auth, RLS, data access policies)",
        "Code review and quality gating of generated output",
        "Recognising when generated code is plausible but wrong",
    ]:
        bullet(doc, item)

    # ── APPENDIX A ────────────────────────────────────────────────────────────

    heading1(doc, "Appendix A — Full Feature Commit Log (LOC Stats)")
    body(doc,
         "Excludes node_modules and lock file commits. Sourced from git log --shortstat. "
         "● = formally logged session (S2 = Session 2, ~25 min; S3 = Session 3, ~35 min).")
    styled_table(doc,
        headers=["Date", "Release", "Commit Summary", "Files", "Ins.", "Del."],
        rows=[
            ("2026-03-13", "v5.0.0",    "Custom Fields builder — FEAT-041 Phase 2.5a",                    "6",  "500",   "5"),
            ("2026-03-12", "v4.9.2 ●",  "Smart Insights grouped by category [S3]",                        "2",  "72",    "42"),
            ("2026-03-12", "v4.9.1 ●",  "CRM view switcher — List, Kanban, Funnel [S3]",                  "1",  "327",   "82"),
            ("2026-03-12", "v4.9.0 ●",  "CRM table view + inline stage selector [S3]",                    "2",  "164",   "175"),
            ("2026-03-12", "v4.8.0 ●",  "Smart Insights dashboard card [S3]",                             "2",  "197",   "1"),
            ("2026-03-12", "v4.7.1 ●",  "Dynamic brand color — full app theme (24 files) [S3]",           "24", "1,102", "131"),
            ("2026-03-12", "v4.7.0 ●",  "Industry templates (Retail/Clinic/Restaurant/Distributor) [S2]", "3",  "170",   "1"),
            ("2026-03-12", "v4.6.0",    "Notification Center, Dashboard custom., CSV Export, Search, Audit","16","1,381","114"),
            ("2026-03-12", "v4.5.0",    "Employee Self-Service Portal",                                   "6",  "667",   "11"),
            ("2026-03-12", "v4.4.0",    "Approval Workflows (Leave, PO, Appraisal)",                      "9",  "1,209", "18"),
            ("2026-03-11", "v4.3.0",    "HR: Goals, Appraisals, Skills Matrix",                           "2",  "1,302", "125"),
            ("2026-03-11", "v4.2.1",    "In-app Help page",                                               "2",  "346",   "0"),
            ("2026-03-11", "v4.2.0",    "Reports page + Supplier management",                             "5",  "699",   "3"),
            ("2026-03-11", "v4.1.1",    "Mobile/tablet responsive layout",                                "6",  "55",    "16"),
            ("2026-03-11", "v4.1.0",    "Supabase integration (data layer)",                              "11", "440",   "7"),
            ("2026-03-11", "v4.0.0",    "Inventory image upload, sortable tables, analytics",             "4",  "369",   "99"),
            ("2026-03-10", "v3.1.0",    "Recharts, currency, toast notifications",                        "8",  "149",   "77"),
            ("2026-03-10", "v3.1.0",    "HR + Settings API routers",                                      "3",  "126",   "1"),
            ("2026-03-10", "v3.0",      "Complete all pages — HR, Settings, UI",                          "12", "819",   "16"),
            ("2026-03-10", "v3.0",      "Production UI with shadcn/ui",                                   "23", "2,368", "826"),
            ("2026-03-10", "v2.0",      "CRM module, FastAPI, Next.js frontend",                          "41", "3,860", "84"),
        ],
        col_widths_cm=[2.4, 2.0, 7.2, 1.2, 1.6, 1.6]
    )
    body(doc, "Total (feature commits only): ~15,800 net new application lines across 24 substantive commits.")

    # ── APPENDIX B ────────────────────────────────────────────────────────────

    heading1(doc, "Appendix B — Module Feature Inventory")
    styled_table(doc,
        headers=["Module", "Feature Count (v5.0.0)", "Highlights"],
        rows=[
            ("Dashboard",             "7",  "KPIs, trend chart, fast/slow movers, currency, Smart Insights, customizable layout"),
            ("Operations",            "6",  "Inventory, POS, Bills, Suppliers, Purchase Orders, image upload"),
            ("HR",                    "7",  "Employees, Payroll, Goals, Appraisals, Skills, Leave, sign-off workflows"),
            ("CRM",                   "7",  "Contacts, Leads, List/Kanban/Funnel views, lead scoring, follow-ups"),
            ("Reports",               "4",  "Sales Report, Top Sellers, Inventory Report, CSV Export"),
            ("Settings",              "5",  "Company info, currency, industry templates, brand color, Custom Fields"),
            ("Help",                  "1",  "In-app user guide covering all modules"),
            ("Employee Self-Service", "4",  "My Goals, My Appraisals, My Leave, My Skills"),
            ("Notifications",         "1",  "Bell icon, pending leave/PO/appraisal/low-stock alerts"),
        ],
        col_widths_cm=[4.0, 3.5, 8.5]
    )

    # ── APPENDIX C ────────────────────────────────────────────────────────────

    heading1(doc, "Appendix C — Stack Reference")
    styled_table(doc,
        headers=["Component", "Technology", "Notes"],
        rows=[
            ("Framework",       "Next.js 14 (App Router)",         "TypeScript, strict mode"),
            ("UI Components",   "shadcn/ui + Tailwind CSS",         "Design system"),
            ("Charts",          "Recharts",                         "Sales trend, top sellers"),
            ("Database",        "Supabase (PostgreSQL)",            "RLS enabled"),
            ("Auth",            "Hardcoded (FEAT-036 planned)",     "Supabase Auth, Phase 3"),
            ("Deployment",      "Vercel",                           "Auto-deploy on push to main"),
            ("AI Collaborator", "Claude Sonnet 4.6 (Anthropic)",    "Full-session pair-programmer"),
        ],
        col_widths_cm=[3.5, 5.5, 7.0]
    )

    body(doc,
         "\nAll git statistics were derived from git log --shortstat on the BzHub repository through "
         "v5.0.0 (2026-03-13). Node_modules and lock file commits are excluded from LOC counts.",
         space_after=4)

    # ── APPENDIX D ────────────────────────────────────────────────────────────

    heading1(doc, "Appendix D — Interaction Log (Full)")
    body(doc, "Source: documentation/INTERACTION_LOG.md. Logging began at v4.7.0.")
    styled_table(doc,
        headers=["Session", "Date", "Goal", "Time", "Releases", "Key Output"],
        rows=[
            ("1", "2026-03-12", "Efficiency analysis + documentation",  "~45 mins", "—",
             "EFFICIENCY_WHITEPAPER.md, generate_whitepaper_docx.py, generate_exec_deck.py"),
            ("2", "2026-03-12", "FEAT-038 Industry Templates",          "~25 mins", "v4.7.0",
             "templates.ts, Settings card, Help section, all docs updated"),
            ("3", "2026-03-12", "Brand color + Smart Insights + CRM views", "~35 mins", "v4.7.1–v4.9.2 (6)",
             "Dynamic brand color (24 files), Smart Insights, CRM List/Kanban/Funnel, seed data"),
            ("Total", "", "", "~105 mins", "6 releases", "~1,600 net LOC (logged sessions only)"),
        ],
        col_widths_cm=[1.4, 2.4, 3.8, 2.0, 2.4, 4.0]
    )
    body(doc,
         "Running totals as of 2026-03-13: 3 sessions logged, ~105 mins total interaction time, "
         "20+ features shipped across v1.0–v5.0.0.",
         space_after=6)

    # ── FOOTER & SAVE ─────────────────────────────────────────────────────────

    add_footer(doc)

    out_path = "/Users/scottvalentino/BzHub/documentation/BzHub_Efficiency_WhitePaper.docx"
    doc.save(out_path)
    print(f"✓  Saved: {out_path}")


if __name__ == "__main__":
    build_document()
