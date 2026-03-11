from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─── Page margins ────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1.1)
section.right_margin  = Inches(1.1)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ─── Colour palette ──────────────────────────────────────────────────────────
PURPLE      = RGBColor(0x6D, 0x28, 0xD9)   # brand purple
DARK_PURPLE = RGBColor(0x4C, 0x1D, 0x95)
DARK        = RGBColor(0x1F, 0x2D, 0x3D)   # near-black
MID         = RGBColor(0x4B, 0x55, 0x63)   # body grey
LIGHT_GREY  = RGBColor(0xF3, 0xF4, 0xF6)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
GREEN       = RGBColor(0x05, 0x96, 0x69)
RED         = RGBColor(0xDC, 0x26, 0x26)
AMBER       = RGBColor(0xD9, 0x77, 0x06)

# ─── Helpers ─────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def cell_borders(cell, color_hex='D1D5DB', sz=4):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color_hex)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def para_spacing(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'),  str(after))
    pPr.append(spacing)

def add_run(para, text, bold=False, italic=False,
            size=11, color=None, font_name='Calibri'):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font_name
    if color:
        run.font.color.rgb = color
    return run

def add_paragraph(text='', bold=False, italic=False,
                  size=11, color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
                  before=60, after=60, font_name='Calibri'):
    p = doc.add_paragraph()
    p.alignment = align
    para_spacing(p, before, after)
    if text:
        add_run(p, text, bold=bold, italic=italic,
                size=size, color=color, font_name=font_name)
    return p

def heading1(text):
    """Large section heading with purple left-bar feel."""
    p = doc.add_paragraph()
    para_spacing(p, before=240, after=80)
    # coloured rule via border on paragraph bottom
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '8')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '6D28D9')
    pBdr.append(bottom)
    pPr.append(pBdr)
    add_run(p, text, bold=True, size=15, color=PURPLE, font_name='Calibri')
    return p

def heading2(text):
    p = doc.add_paragraph()
    para_spacing(p, before=200, after=60)
    add_run(p, text, bold=True, size=12.5, color=DARK_PURPLE, font_name='Calibri')
    return p

def heading3(text):
    p = doc.add_paragraph()
    para_spacing(p, before=160, after=40)
    add_run(p, text, bold=True, size=11.5, color=DARK, font_name='Calibri')
    return p

def body(text, before=40, after=40):
    return add_paragraph(text, size=11, color=MID,
                         before=before, after=after)

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    para_spacing(p, before=20, after=20)
    indent_val = str(360 + level * 360)
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), indent_val)
    pPr.append(ind)
    add_run(p, text, size=11, color=MID, font_name='Calibri')
    return p

def tick(text, color=GREEN):
    p = doc.add_paragraph(style='List Bullet')
    para_spacing(p, before=20, after=20)
    add_run(p, '✅  ' if color == GREEN else ('❌  ' if color == RED else '🔮  '),
            size=11, font_name='Segoe UI Emoji')
    add_run(p, text, size=11, color=MID, font_name='Calibri')
    return p

def add_table(headers, rows,
              header_bg='6D28D9', header_fg=WHITE,
              alt_bg='F5F3FF', border_color='D1D5DB',
              col_widths=None):
    """Styled table with coloured header and alternating row shading."""
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'

    # header row
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr_cells[i], header_bg)
        cell_borders(hdr_cells[i], border_color)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_spacing(p, before=60, after=60)
        add_run(p, h, bold=True, size=10, color=header_fg)

    # data rows
    for ri, row_data in enumerate(rows):
        bg = alt_bg if ri % 2 == 0 else 'FFFFFF'
        cells = t.rows[ri+1].cells
        for ci, val in enumerate(row_data):
            set_cell_bg(cells[ci], bg)
            cell_borders(cells[ci], border_color)
            p = cells[ci].paragraphs[0]
            para_spacing(p, before=50, after=50)
            # allow inline colour markers: use tuple (text, RGBColor)
            if isinstance(val, tuple):
                add_run(p, val[0], size=10, color=val[1])
            else:
                add_run(p, str(val), size=10, color=MID)

    # column widths
    if col_widths:
        for ri in range(len(t.rows)):
            for ci, w in enumerate(col_widths):
                t.rows[ri].cells[ci].width = Inches(w)
    return t

def page_break():
    doc.add_page_break()

def divider():
    p = doc.add_paragraph()
    para_spacing(p, before=120, after=120)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), 'E5E7EB')
    pBdr.append(bot)
    pPr.append(pBdr)

# ═══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
para_spacing(p, before=800, after=0)
add_run(p, 'BzHub', bold=True, size=42, color=PURPLE, font_name='Calibri')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
para_spacing(p, before=60, after=0)
add_run(p, 'Business Operating System for SMEs', bold=False,
        size=18, color=DARK_PURPLE, font_name='Calibri')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

divider()

p = doc.add_paragraph()
para_spacing(p, before=40, after=40)
add_run(p, 'From Idea to Execution — The Full Story, Honest Assessment & Path Forward',
        italic=True, size=13, color=MID, font_name='Calibri')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

for line in [
    '',
    'Prepared by:  Scott Valentino',
    'Document Date:  March 10, 2026',
    'Version:  1.0  — Comprehensive Business Overview',
]:
    p = doc.add_paragraph()
    para_spacing(p, before=20, after=20)
    add_run(p, line, size=11, color=MID)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════════════════════════
heading1('Table of Contents')

toc_items = [
    ('1.', 'Executive Summary'),
    ('2.', 'The Problem We Are Solving'),
    ('3.', 'The Idea — Inception'),
    ('4.', 'The Journey — What Was Built'),
    ('5.', 'Where We Are Today'),
    ('6.', 'Design Decisions — The Good, The Bad, The Honest'),
    ('7.', 'The Market'),
    ('8.', 'The Way Forward — Roadmap'),
    ('9.', 'The Pitch'),
    ('10.', 'Financial Model (Indicative)'),
    ('11.', 'Risks & Mitigation'),
    ('12.', 'Summary'),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    para_spacing(p, before=30, after=30)
    add_run(p, f'{num}  ', bold=True, size=11, color=PURPLE)
    add_run(p, title, size=11, color=MID)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════════════════
heading1('1.  Executive Summary')

body('BzHub is a modular, all-in-one business management platform built for small and medium enterprises (SMEs). It consolidates the fragmented tooling that kills small business productivity — inventory management, point-of-sale, CRM, HR, payroll, analytics, and reporting — into a single, cohesive system.')

body('The current state: a fully working desktop application (Python/Tkinter), a REST API backend (FastAPI), and the early shell of a web frontend (Next.js). The architecture has been carefully designed to scale from a single-user desktop app to a cloud-hosted, multi-tenant SaaS platform — without a rebuild.')

body('The insight that drives it: every small business owner is using at least 4–6 disconnected tools. They pay for each. They manually reconcile data between them. They lose hours every week. BzHub makes that stop.')

body('Where it is going: web-first, cloud-hosted SaaS. Subscription pricing. Indian SME market as the initial beachhead, global expansion as the second phase.')

divider()

# Key metrics callout table (3-column highlight strip)
t = doc.add_table(rows=1, cols=3)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
labels = ['Modules Built', 'Tests Passing', 'Architecture Layers']
values = ['12+', '24 / 100%', '3 (Core → Service → UI)']
for ci, (lbl, val) in enumerate(zip(labels, values)):
    set_cell_bg(t.rows[0].cells[ci], '6D28D9')
    cell_borders(t.rows[0].cells[ci], '4C1D95', sz=6)
    p = t.rows[0].cells[ci].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(p, before=120, after=40)
    add_run(p, val + '\n', bold=True, size=16, color=WHITE)
    add_run(p, lbl, size=9, color=RGBColor(0xC4, 0xB5, 0xFD))
for ci in range(3):
    t.rows[0].cells[ci].width = Inches(2.0)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  2. THE PROBLEM
# ═══════════════════════════════════════════════════════════════════════════════
heading1('2.  The Problem We Are Solving')

body('Walk into any small business — a retail shop, a service center, a small clinic, a trading company with 10 employees. Ask the owner what software they use. You will hear a version of this:')

p = doc.add_paragraph()
para_spacing(p, before=80, after=80)
pPr = p._p.get_or_add_pPr()
ind = OxmlElement('w:ind'); ind.set(qn('w:left'), '720'); pPr.append(ind)
pBdr = OxmlElement('w:pBdr')
left = OxmlElement('w:left')
left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '12')
left.set(qn('w:space'), '10'); left.set(qn('w:color'), '6D28D9')
pBdr.append(left); pPr.append(pBdr)
add_run(p, '"We have Tally for accounts, Excel for inventory, WhatsApp to communicate with customers, a different app for billing, and we track HR stuff manually."',
        italic=True, size=11.5, color=DARK_PURPLE)

body('This is the status quo for the vast majority of small businesses in India and across developing markets. The tools are disconnected. The data lives in silos. The business owner is the human middleware — manually copying numbers from one place to another, spending mental energy on administrative overhead instead of growth.')

heading2('The Specific Pain Points')

add_table(
    headers=['Pain Point', 'What Is Happening Today'],
    rows=[
        ('Inventory gaps', 'Stock runs out without warning. No automated alerts. Tracked in Excel.'),
        ('Sales reconciliation', 'POS data in one app, inventory in another — never in sync.'),
        ('Customer tracking', 'Contacts in WhatsApp, leads in a spreadsheet, no pipeline visibility.'),
        ('HR fragmentation', 'Employee records in folders, payroll in Excel, appraisals not done.'),
        ('Reporting blindness', 'No real-time dashboards. Reports are last month\'s data, manually compiled.'),
        ('Tool cost accumulation', '5 SaaS tools at $30/month each = $150+/month for a business with $3,000 revenue.'),
    ],
    col_widths=[1.9, 4.2],
)

body('')
heading2('Who Suffers Most')
body('Small businesses in the 2–50 employee range — the "missing middle" of software. Enterprise ERP (SAP, Oracle) is too expensive and complex. Consumer apps are too simple. This bracket is underserved by design.')

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  3. INCEPTION
# ═══════════════════════════════════════════════════════════════════════════════
heading1('3.  The Idea — Inception')

body('BzHub started as a practical frustration, not a pitch-deck idea.')
body('The original concept was simple: one app that a shop owner can run on a laptop, that handles their inventory, their billing, and their daily sales report. No cloud required. No subscription. Just download and run.')
body('The early technical choice reflected that simplicity: Python + SQLite. No external servers, no network dependency, no configuration headache.')
body('But as modules were added — CRM, HR, payroll, appraisals, analytics — the product grew into something more ambitious. What started as a POS system became an ERP. And the realization emerged: the architecture needed to match the ambition.')

heading2('The Founding Insight')
body('Most ERP systems are built desktop-first or cloud-first. Rarely both. BzHub\'s founding insight was that the business logic should live in exactly one place — a service layer — and the delivery mechanism (desktop, web browser, mobile) should be interchangeable. Build once. Deploy anywhere.')
body('This sounds obvious in hindsight. It was not obvious in the original monolithic codebase. It required a deliberate refactoring to achieve.')

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  4. THE JOURNEY
# ═══════════════════════════════════════════════════════════════════════════════
heading1('4.  The Journey — What Was Built')

heading2('Phase 1 — The Monolith  (Early 2026)')
body('BzHub began as a single Python file that grew into a 3,000+ line monolith. Everything was in one place: database calls, business logic, UI rendering. It worked, but it was impossible to test, extend safely, or separate from the UI layer.')

heading3('What was built in Phase 1:')
for item in [
    'Inventory management (CRUD, search, low-stock alerts)',
    'Point of Sale — cart, checkout, tax/discount, receipt printing',
    'Bills / sales history with date filtering',
    'Visitor contact management',
    'Basic dashboard with KPI cards',
    'Employee HR management and payroll tracking',
    'Settings and SMTP email alerts',
    'Login and authentication',
]:
    bullet(item)

body('The problem: every new feature risked breaking everything else. There were no tests. Business logic was entangled with UI code. The database was called directly from the UI layer.')

divider()

heading2('Phase 2 — The Refactoring  (v1.0, February 2026)')
body('This was the turning point. The monolith was dismantled and replaced with a clean three-layer architecture:')

# Architecture diagram table
t = doc.add_table(rows=3, cols=1)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
layers = [
    ('UI Layer  (Desktop / Web / Mobile)', 'F5F3FF', DARK_PURPLE),
    ('Service Layer  (Business Logic)', 'EDE9FE', DARK_PURPLE),
    ('Database Layer  (Swappable Adapter)', '6D28D9', WHITE),
]
for ri, (label, bg, fg) in enumerate(layers):
    set_cell_bg(t.rows[ri].cells[0], bg.replace('#',''))
    cell_borders(t.rows[ri].cells[0], '6D28D9', sz=8)
    p = t.rows[ri].cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(p, before=100, after=100)
    add_run(p, label, bold=True, size=11, color=fg)
    t.rows[ri].cells[0].width = Inches(5.0)

body('')
body('What the refactoring achieved:')
add_table(
    headers=['Before', 'After'],
    rows=[
        ('3,122-line monolith', '12 modular tab files (~2,800 total lines)'),
        ('No tests', '24 unit + integration tests, 100% pass rate'),
        ('No logging', 'Python logging module, structured log output'),
        ('Hardcoded admin credentials', 'Environment variable–based configuration'),
        ('Direct DB calls from UI', 'Service layer as the sole DB consumer'),
        ('No database abstraction', 'Abstract DatabaseAdapter interface'),
        ('No documentation', '21 markdown reference documents'),
    ],
    col_widths=[2.9, 3.2],
)

divider()

heading2('Phase 3 — CRM + API + Web  (v2.0, March 9, 2026)')
body('With a stable architecture, v2.0 added three major pillars without touching existing features.')

heading3('CRM Module')
for item in [
    'Contacts directory with full CRUD',
    '6-stage sales pipeline: New → Contacted → Qualified → Proposal → Won → Lost',
    'Lead cards with value, probability, and owner assignment',
    'Activity log per lead (calls, emails, meetings, notes)',
    'Pipeline analytics: conversion rate and total pipeline value',
]:
    bullet(item)

heading3('REST API  (FastAPI)')
for item in [
    '6 routers: Auth, Inventory, Sales, Contacts, Leads, Dashboard',
    'Auto-generated interactive docs at /docs',
    'CORS enabled for web frontend consumption',
    'Shared dependency injection for services',
]:
    bullet(item)

heading3('Web Frontend  (Next.js)')
for item in [
    'Login page with JWT authentication flow',
    'Dashboard page with live KPI cards',
    'Operations hub page',
    'CRM Kanban pipeline page',
    'Tailwind CSS 4 with purple brand design system',
    'shadcn/ui component library integrated',
]:
    bullet(item)

heading3('3 Critical Bugs Fixed')
add_table(
    headers=['Bug', 'Fix Applied'],
    rows=[
        ('Hardcoded admin credentials in source', 'Moved to environment variables — ADMIN_USERNAME / ADMIN_PASSWORD'),
        ('print() used throughout for debugging', 'Replaced with Python logging module, written to bizhub.log'),
        ('No database indexes on query columns', '7 indexes added — 10–100x faster queries on large datasets'),
    ],
    col_widths=[2.4, 3.7],
)

divider()

heading2('Phase 4 — Modernization Decision  (March 10, 2026)')
body('A deliberate strategic decision was made: the web UI is the future, not the desktop.')

add_table(
    headers=['Option', 'Approach', 'Verdict'],
    rows=[
        ('Option A', 'Upgrade Tkinter with customtkinter', ('Rejected — Tkinter ceiling remains', RED)),
        ('Option B', 'Build out web UI with shadcn/ui', ('Chosen — no design ceiling, any device', GREEN)),
    ],
    col_widths=[1.0, 2.5, 2.6],
)

body('')
body('The desktop app remains functional for backward compatibility. All new development goes into the web frontend. This decision positions BzHub for SaaS delivery.')

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  5. WHERE WE ARE TODAY
# ═══════════════════════════════════════════════════════════════════════════════
heading1('5.  Where We Are Today')

heading2('Feature Status')

add_table(
    headers=['Module', 'Status', 'Platform'],
    rows=[
        ('Inventory Management',      ('✅ Done', GREEN),    'Desktop'),
        ('Point of Sale',             ('✅ Done', GREEN),    'Desktop'),
        ('CRM Contacts',              ('✅ Done', GREEN),    'Desktop + API'),
        ('CRM Pipeline (Kanban)',      ('✅ Done', GREEN),    'Desktop + API'),
        ('HR — Employees',            ('✅ Done', GREEN),    'Desktop'),
        ('HR — Payroll',              ('✅ Done', GREEN),    'Desktop'),
        ('HR — Appraisals (360°)',    ('✅ Done', GREEN),    'Desktop'),
        ('Dashboard & Analytics',     ('✅ Done', GREEN),    'Desktop + API'),
        ('Reports (period selector)', ('✅ Done', GREEN),    'Desktop'),
        ('Visitor Log',               ('✅ Done', GREEN),    'Desktop'),
        ('Email Alerts (SMTP)',        ('✅ Done', GREEN),    'Desktop'),
        ('Excel Import / Export',     ('✅ Done', GREEN),    'Desktop'),
        ('Activity Logging',          ('✅ Done', GREEN),    'Desktop'),
        ('Dark Mode',                 ('✅ Done', GREEN),    'Desktop'),
        ('REST API (6 routers)',       ('✅ Done', GREEN),    'FastAPI'),
        ('Web — Login',               ('🔄 Skeleton', AMBER), 'Next.js'),
        ('Web — Dashboard',           ('🔄 Skeleton', AMBER), 'Next.js'),
        ('Web — Operations',          ('🔄 Skeleton', AMBER), 'Next.js'),
        ('Web — CRM',                 ('🔄 Skeleton', AMBER), 'Next.js'),
        ('PostgreSQL Support',        ('🔮 Planned', MID),   '—'),
        ('Multi-tenant',              ('🔮 Planned', MID),   '—'),
        ('Payment Gateway',           ('🔮 Planned', MID),   '—'),
    ],
    col_widths=[2.4, 1.4, 2.3],
)

body('')
heading2('Tech Stack')

add_table(
    headers=['Layer', 'Technology'],
    rows=[
        ('Backend',      'Python 3.x, FastAPI, SQLite (→ PostgreSQL via adapter)'),
        ('Desktop UI',   'Tkinter, matplotlib'),
        ('Web UI',       'Next.js 14, TypeScript, Tailwind CSS 4, shadcn/ui, Recharts'),
        ('Testing',      'pytest — 24 tests, 100% pass rate'),
        ('Database',     'SQLite (local), DatabaseAdapter ready for PostgreSQL'),
        ('Documentation','21 markdown files — architecture, bugs, release notes, quickstart'),
    ],
    col_widths=[1.5, 4.6],
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  6. DESIGN DECISIONS
# ═══════════════════════════════════════════════════════════════════════════════
heading1('6.  Design Decisions — The Good, The Bad, The Honest')

body('This section is written with full honesty. Every product has decisions they got right and decisions they would revisit. Here are both.')

# ── GOOD ────────────────────────────────────────────────────────────────────
heading2('The Good Decisions')

heading3('✅  1.  Service Layer Architecture')
body('All business logic lives in service classes (InventoryService, CRMService, HRService). UI never touches the database directly.')
body('Why it was right: When FastAPI was added, there was zero code duplication. The API routers call the same service methods the desktop uses. When the web frontend arrives, same story. One source of truth.')
body('Impact: Adding a mobile app takes hours, not weeks. The services are already there.')

heading3('✅  2.  Database Abstraction Layer')
body('An abstract DatabaseAdapter interface defines all database operations. SQLiteAdapter is the first implementation. PostgreSQLAdapter will be next.')
body('Why it was right: SQLite is perfect for local desktop use — zero config, single file, fast. But it cannot serve concurrent users in the cloud. The abstraction means BzHub can migrate without rewriting anything above the adapter layer. This is worth months of future work already done.')

heading3('✅  3.  FastAPI Choice')
body('FastAPI was chosen over Flask for the REST API.')
body('Why it was right: Async support, Pydantic validation, and auto-generated interactive API docs at /docs out of the box. The docs endpoint is also a practical demo tool — show it to a developer partner and they can explore the API immediately.')

heading3('✅  4.  shadcn/ui for the Web Component Library')
body('shadcn/ui (Radix UI + Tailwind) was chosen over Material-UI, Chakra UI, or custom components.')
body('Why it was right: It is what Vercel, Linear, and modern B2B SaaS companies use. It owns no global styles, is copy-paste composable, fully accessible via Radix primitives, and produces a professional UI that competes visually with any modern SaaS product. There is no design ceiling.')

heading3('✅  5.  Writing Tests Before Adding More Features')
body('At the v1.0 refactor, 24 tests were written before v2.0 feature development began.')
body('Why it was right: The CRM module and API were added without breaking any existing functionality. The tests were the safety net. Without them, each new feature would have been a gamble.')

heading3('✅  6.  Documenting Architecture Decisions')
body('A /documentation/ folder was maintained with 21 markdown files covering architecture, refactoring history, bug fixes, release notes, and code review recommendations.')
body('Why it was right: Onboarding a developer, technical co-founder, or an investor\'s technical advisor becomes significantly easier when the architectural reasoning is written down. This document exists because that work was done.')

divider()

# ── BAD ─────────────────────────────────────────────────────────────────────
heading2('The Bad Decisions  (Honest)')

heading3('❌  1.  Starting With a Monolith and No Tests')
body('Why it was wrong: By the time the codebase was 3,000+ lines, refactoring was a major project rather than a routine activity. A service layer and test suite from day one would have made every subsequent feature addition clean.')
body('Lesson: Structure is not a future problem. It is a day-one decision.')

heading3('❌  2.  Tkinter as the UI Framework')
body('Why it was understandable but limiting: Tkinter is 1990s-era GUI technology. It produces UIs that look dated regardless of theming effort. Responsive layouts are nearly impossible. Accessibility is poor.')
body('The honest impact: When a business owner or investor opens a Tkinter app today, they compare it to modern web apps. The comparison is unfavorable — not because the features are weak, but because the visual language feels outdated.')
body('What was done: The strategic decision was made to pivot to the web UI. Tkinter is maintained for offline users but is not the product\'s future face. This was the right correction.')

heading3('❌  3.  Hardcoded Admin Credentials in Source Code')
body('Why it was wrong: This is a security vulnerability by definition. Any developer with read access to the repo has working credentials. Any user who installs the app without changing defaults is exposed.')
body('What was done: Fixed in v2.0 — credentials now read from environment variables. But this should never have been in the codebase.')
body('Lesson: Default-secure is the only acceptable baseline.')

heading3('❌  4.  Using print() for All Debugging')
body('Why it was wrong: print() provides no log levels, no timestamps, no file output, no filtering, and no structured format. In production, these become noise that cannot be turned off.')
body('What was done: Replaced throughout with Python\'s logging module in v2.0.')

heading3('❌  5.  No Database Indexes on Frequently-Queried Columns')
body('Why it was wrong: SQLite queries on unindexed columns are full table scans. With 1,000+ inventory rows and 10,000+ sales records, query times degrade noticeably. Indexes are not premature optimisation — they are table stakes.')
body('What was done: 7 indexes added in v2.0. The fix was one session of work. The omission created unnecessary risk.')

heading3('❌  6.  SQLite Is Not Suitable for Multi-User Cloud Deployment')
body('Why it is a current gap: SQLite does not support concurrent writes. A web app with two users making changes simultaneously will have race conditions.')
body('What is needed: A PostgreSQLAdapter that implements the same DatabaseAdapter interface. The architecture is ready. The implementation has not been written yet. This is the most pressing technical debt for cloud launch.')

heading3('❌  7.  The Web Frontend Is Still a Skeleton')
body('Why this creates a credibility gap: The desktop app is fully functional. The API is fully functional. But the web frontend — which is the product\'s future face — has skeleton pages with no real interactivity. A potential user or investor clicking through the web app today will see incomplete work.')
body('What is needed: The web frontend must be brought to feature parity with the desktop app. This is the immediate priority.')

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  7. THE MARKET
# ═══════════════════════════════════════════════════════════════════════════════
heading1('7.  The Market')

heading2('Market Size')

add_table(
    headers=['Segment', 'Definition', 'Size / Value'],
    rows=[
        ('TAM — Total Addressable Market',
         'All SMEs globally that need business management software',
         '$1.2 trillion global SMB software market by 2030 (IDC)'),
        ('SAM — Serviceable Addressable Market',
         'India-based SMEs, 2–50 employees, digitally-ready sectors',
         '~5 million businesses / ~$1–5B ARR potential'),
        ('SOM — Serviceable Obtainable Market (Yr 1–2)',
         'Early-adopter SMEs reachable via digital channels',
         '500–2,000 customers / $108K–$1.4M ARR'),
    ],
    col_widths=[1.7, 2.8, 1.8],
)

body('')
heading2('Competitive Landscape')

add_table(
    headers=['Competitor', 'Strength', 'Weakness', 'BzHub Advantage'],
    rows=[
        ('Tally ERP',   'Trusted brand, accountant-familiar',  'Steep learning curve, no web UI',     'Modern UX, web-native, easier to use'),
        ('Zoho Books',  'Feature-rich, good integrations',     'Complex, expensive for micro-SME',    'Simpler, all-in-one, lower price'),
        ('QuickBooks',  'Strong in accounting',                'Not a full ERP, US-focused',          'Full ERP, India-ready, CRM + HR'),
        ('Vyapar',      'India-focused, GST-compliant',        'Billing only — no CRM or HR',         'Full suite including CRM and HR'),
        ('Odoo',        'Open source, very modular',           'Extremely complex to set up',         'Lightweight, fast setup, opinionated'),
        ('Spreadsheets','Free, familiar',                      'Not scalable, manual, error-prone',   'Automated, structured, interconnected'),
    ],
    col_widths=[1.2, 1.6, 1.7, 1.7],
)

body('')
body('BzHub\'s position: the only platform in the India SME market that combines Inventory + POS + CRM pipeline + HR + Payroll + Analytics in a single product, at a price point accessible to businesses under ₹1 crore annual revenue, with no technical expertise required to deploy.')

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  8. THE ROADMAP
# ═══════════════════════════════════════════════════════════════════════════════
heading1('8.  The Way Forward — Roadmap')

heading2('Immediate Priority  (Next 30 Days) — Make It Shippable')
body('The single most important near-term milestone is a complete, functional web app that matches the desktop app\'s capabilities.')

heading3('Sprint 1 — Web Frontend Completion')
for item in [
    'Build Inventory management page with shadcn/ui DataTable',
    'Build POS page with cart and checkout flow',
    'Build CRM Contacts page with CRUD modals',
    'Build CRM Pipeline as interactive Kanban board (drag-and-drop)',
    'Build HR — Employees page',
    'Build Reports page with Recharts visualisations',
    'Add toast notifications, loading states, and error handling throughout',
    'Responsive layout for tablet and mobile',
]:
    bullet(item)

heading3('Sprint 2 — Database for Cloud')
for item in [
    'Implement PostgreSQLAdapter (same interface as SQLiteAdapter)',
    'Set up connection pooling for concurrent users',
    'Environment-based database switching (SQLite for dev, PostgreSQL for prod)',
    'Write migration scripts for initial schema',
]:
    bullet(item)

heading3('Sprint 3 — Deployment')
for item in [
    'Dockerise the FastAPI backend',
    'Deploy to Railway / Render / AWS (initial)',
    'Configure environment variables for production',
    'Set up HTTPS, custom domain, rate limiting, and security headers',
]:
    bullet(item)

divider()

heading2('3–6 Month Horizon — Beta to Paying Customers')

heading3('User Management')
for item in [
    'Multi-user roles (Admin, Manager, Staff, Read-only)',
    'Invite system for team members',
    'Audit trail for all user actions',
    'Session management and JWT refresh tokens',
]:
    bullet(item)

heading3('Product Features')
for item in [
    'Supplier management (purchase orders, goods receipt)',
    'Customer loyalty programme (points, tiers)',
    'Invoice generation and PDF export',
    'GST compliance (India) — GSTIN on invoices, tax reports',
    'SMS / WhatsApp alerts for low stock and order confirmations',
]:
    bullet(item)

heading3('Commercial')
for item in [
    'Subscription billing integration (Razorpay or Stripe)',
    'Free trial — 14 days, no credit card required',
    'Pricing tiers: Starter / Growth / Business',
    'Customer onboarding flow (guided setup wizard)',
]:
    bullet(item)

divider()

heading2('6–18 Month Horizon — Scale')

heading3('Platform Expansion')
for item in [
    'Mobile PWA (Progressive Web App) — works on any smartphone browser',
    'Tally XML import for migrating businesses',
    'Bank statement reconciliation import',
    'Inventory forecasting (ML-based reorder suggestions)',
    'API marketplace for third-party integrations',
]:
    bullet(item)

heading3('Business Expansion')
for item in [
    'Channel sales partnerships with CA firms and Tally resellers',
    'Reseller / white-label programme for accounting software firms',
    'WhatsApp Business API for customer notifications',
    'Expansion to Southeast Asia (Indonesia, Philippines)',
]:
    bullet(item)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  9. THE PITCH
# ═══════════════════════════════════════════════════════════════════════════════
heading1('9.  The Pitch')

heading2('One-Liner')

p = doc.add_paragraph()
para_spacing(p, before=100, after=100)
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
left = OxmlElement('w:left')
left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '16')
left.set(qn('w:space'), '12'); left.set(qn('w:color'), '6D28D9')
pBdr.append(left); pPr.append(pBdr)
ind = OxmlElement('w:ind'); ind.set(qn('w:left'), '480'); pPr.append(ind)
add_run(p, 'BzHub is the all-in-one business operating system for small businesses — inventory, billing, CRM, and HR in a single platform, starting at ₹1,500/month.',
        bold=True, size=13, color=DARK_PURPLE, font_name='Calibri')

heading2('The Pitch Story  (2 Minutes)')

heading3('Open with the problem:')
p = doc.add_paragraph()
para_spacing(p, before=60, after=60)
add_run(p, '"A shop owner I know runs a retail electronics store in Bangalore. He uses Tally for accounts, Excel for inventory, WhatsApp to follow up with customers, and a separate billing app on his counter. Every Friday, he sits down for three hours to reconcile these systems manually. He is not doing business on Friday afternoons. He is doing data entry."',
        italic=True, size=11, color=MID)

heading3('The opportunity:')
body('"India has 63 million small businesses. Most of them operate exactly like this. The gap between \'Excel\' and \'SAP\' has never been properly filled. No one has built the default business operating system for the emerging market SME."')

heading3('The product:')
body('"BzHub is a single platform that replaces all of those tools. One login. Inventory updates automatically when you make a sale. Your CRM knows which customers have not heard from you in 30 days. Your dashboard shows this week\'s revenue versus last week\'s at a glance."')

heading3('The architecture — why we can win:')
body('"We did not build a feature — we built a platform. The business logic is written once and delivered anywhere: desktop, browser, mobile. We can switch our database from SQLite to PostgreSQL for cloud scale without touching the business logic. We have 24 automated tests. We already have a REST API."')

heading3('The ask:')
body('"We are raising ₹[X] to complete the web frontend, deploy to cloud, and acquire our first 100 paying customers. With that investment, we reach cash-flow break-even at 250 customers — a target we can achieve in 12 months."')

divider()

heading2('Pitch Deck Slide Structure')

add_table(
    headers=['Slide', 'Content'],
    rows=[
        ('1',  'Title — BzHub, The Business Operating System for SMEs'),
        ('2',  'The Problem — 5 disconnected tools, manual reconciliation, 3 hours/week lost'),
        ('3',  'The Market — 63M SMEs in India, $1.2T global SMB software market'),
        ('4',  'The Product — Screen recordings of dashboard, CRM, inventory in action'),
        ('5',  'How It Works — 3-layer architecture diagram, desktop + web + API'),
        ('6',  'The Competition — Why BzHub beats Tally, Zoho, Vyapar in this segment'),
        ('7',  'Business Model — Subscription tiers, pricing, LTV/CAC assumptions'),
        ('8',  'Traction — Current features, tests, API, architecture discipline'),
        ('9',  'The Roadmap — 30 days / 6 months / 18 months'),
        ('10', 'The Team — Background and credibility'),
        ('11', 'The Ask — Amount raised, use of funds, milestones unlocked'),
        ('12', 'The Vision — Default business OS for emerging market SMEs'),
    ],
    col_widths=[0.5, 5.6],
)

body('')
heading2('Technical Proof Points  (for investor Q&A)')

body('These details answer "have you really thought this through?" before it is asked:')
for item in [
    'Modular architecture — 12 independent modules. Proven by CRM addition in v2.0 without breaking existing features.',
    'Database abstraction — SQLite today, PostgreSQL tomorrow. Zero code change above the adapter layer.',
    'Test suite — 24 tests, 100% pass rate. Not "we test manually."',
    'Security baseline — credentials from environment variables. No hardcoded secrets. Parameterised SQL throughout.',
    'REST API — FastAPI with auto-generated docs. Any developer can integrate in an afternoon.',
    'Performance — 7 database indexes in place. Not an afterthought.',
    'Documentation — 21 markdown files including architecture whitepaper. Unusual discipline for an early-stage project.',
]:
    bullet(item)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  10. FINANCIAL MODEL
# ═══════════════════════════════════════════════════════════════════════════════
heading1('10.  Financial Model  (Indicative)')

heading2('Pricing Tiers')

add_table(
    headers=['Tier', 'Price (Monthly)', 'Users', 'Target Customer'],
    rows=[
        ('Starter',    '₹1,499',   'Up to 3',         'Solo operator, micro business'),
        ('Growth',     '₹3,499',   'Up to 10',        '5–15 employee SME'),
        ('Business',   '₹6,999',   'Up to 30',        '15–50 employee company'),
        ('Enterprise', 'Custom',   'Unlimited',       'Large SME, custom integrations'),
    ],
    col_widths=[1.3, 1.4, 1.2, 2.5],
)

body('')
heading2('Revenue Milestones')

add_table(
    headers=['Milestone', 'Customers', 'Mix Assumed', 'MRR', 'ARR'],
    rows=[
        ('Beta launch',      '20',    '80% Starter, 20% Growth',           '₹33,380',    '₹4.0L'),
        ('Break-even',       '250',   '60% Starter, 30% Growth, 10% Biz',  '₹6.6L',      '₹79.2L'),
        ('Year 1 target',    '500',   '50% / 35% / 15% mix',               '₹13.7L',     '₹1.64Cr'),
        ('Year 2 target',    '2,000', '40% / 40% / 20% mix',               '₹60.2L',     '₹7.2Cr'),
    ],
    col_widths=[1.4, 1.1, 2.3, 1.1, 0.9],
)

body('')
heading2('Key Unit Economics')

add_table(
    headers=['Metric', 'Estimate', 'Notes'],
    rows=[
        ('CAC',             '₹3,000–₹8,000',          'Digital + content + partner channel'),
        ('LTV',             '₹36,000–₹84,000',         '24-month retention at Starter/Growth price'),
        ('LTV : CAC ratio', '5:1 to 10:1',             'Healthy SaaS benchmark is 3:1'),
        ('Gross Margin',    '~80%',                     'Software-only, minimal COGS'),
        ('Churn Target',    '<3% monthly',              'SMB industry average is 3–5%'),
    ],
    col_widths=[1.7, 1.8, 2.6],
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  11. RISKS
# ═══════════════════════════════════════════════════════════════════════════════
heading1('11.  Risks & Mitigation')

add_table(
    headers=['Risk', 'Likelihood', 'Impact', 'Mitigation'],
    rows=[
        ('Web frontend takes longer than planned',
         ('Medium', AMBER), ('High', RED),
         'Skeleton already built. Focus all resources here first.'),
        ('PostgreSQL migration breaks something',
         ('Low', GREEN), ('High', RED),
         'DB abstraction limits blast radius. Tests will catch regressions.'),
        ('SMBs resist switching from Tally',
         ('High', RED), ('High', RED),
         'Target Tally non-users first. Offer Tally import tool.'),
        ('Competitive response from Zoho / Vyapar',
         ('Medium', AMBER), ('Medium', AMBER),
         'Focus on underserved micro-SME at accessible price point.'),
        ('Solo developer bottleneck',
         ('High', RED), ('High', RED),
         'Everything is documented. Hire / partner after first revenue.'),
        ('GST compliance gap (India)',
         ('Medium', AMBER), ('High', RED),
         'GST integration is on the 3–6 month roadmap.'),
        ('Security breach in multi-tenant cloud',
         ('Low', GREEN), ('Very High', RED),
         'Tenant isolation, penetration testing, proper JWT/OAuth before launch.'),
    ],
    col_widths=[1.7, 1.0, 0.9, 2.6],
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  12. SUMMARY
# ═══════════════════════════════════════════════════════════════════════════════
heading1('12.  Summary')

body('BzHub is not a side project that accidentally became an ERP. It is a deliberately architected platform that started with a real problem, made pragmatic early decisions, corrected its mistakes with discipline, and is now positioned to make the transition from developer tool to commercial product.')

heading2('The Honest State of Things')

for item in [
    'The product works. The desktop app is full-featured and functional.',
    'The architecture is sound. Service layer, database abstraction, REST API, and test suite are the foundation for a scalable business.',
    'The web frontend is the gap. Everything else being ready makes this a highly leveraged investment of the next 30 days.',
    'The mistakes were made and corrected — that is what matters. Monolith split. Tests written. Security fixed. Indexes added. The codebase reflects learning, not stubbornness.',
]:
    bullet(item)

heading2('The Way Forward Is Clear')

add_table(
    headers=['Step', 'Action', 'Timeline'],
    rows=[
        ('1', 'Complete the web frontend',                      '30 days'),
        ('2', 'Deploy to cloud with PostgreSQL',                '60 days'),
        ('3', 'Acquire first 20 beta customers, collect feedback', '90 days'),
        ('4', 'Price, iterate, and scale',                      'Ongoing'),
    ],
    col_widths=[0.5, 3.8, 1.8],
)

body('')
body('The market is large. The problem is real. The architecture is ready. The product needs to be in front of users.')

divider()

p = doc.add_paragraph()
para_spacing(p, before=160, after=60)
add_run(p,
    'For technical architecture detail, see /documentation/ARCHITECTURE.md\n'
    'For feature history, see /documentation/RELEASE_NOTES_v2.0.md\n'
    'For bug fix history, see /documentation/BUG_FIX_PROGRESS.md',
    italic=True, size=9.5, color=MID)

# ─── Save ────────────────────────────────────────────────────────────────────
output_path = '/Users/scottvalentino/C_Love_Coding/documentation/BzHub_Business_Document.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
